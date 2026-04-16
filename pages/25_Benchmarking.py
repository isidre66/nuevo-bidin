import streamlit as st
from control_acceso import verificar_acceso
verificar_acceso('manager')
from asistentes import mostrar_felix
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import json, os
from datetime import date
from supabase import create_client

def _cargar_indices():
    try:
        url = st.secrets.get("SUPABASE_URL", os.environ.get("SUPABASE_URL",""))
        key = st.secrets.get("SUPABASE_KEY", os.environ.get("SUPABASE_KEY",""))
        ec  = st.session_state.get('empresa_codigo','')
        if not (url and key and ec):
            return
        sb = create_client(url, key)
        emp = sb.table('empresas').select('*').eq('codigo', ec).execute().data or []
        if emp:
            e = emp[0]
            ventas    = float(e.get('ventas') or 0)
            empleados = int(e.get('empleados') or 1)
            roa       = float(e.get('roa') or 0)
            var_vtas  = float(e.get('var_ventas') or 0)
            var_empl  = float(e.get('var_empleados') or 0)
            productiv = ventas / max(empleados, 1)
            endeud    = float(e.get('endeudamiento') or 0)
            exp_cod   = st.session_state.get('save_export_cod', 0)
            ice = min(100, max(0, round(min(roa/20,1)*30 + min(var_vtas/100,1)*25 + min(productiv/500000,1)*25 + (1-min(endeud,1))*20, 1)))
            isf = min(100, max(0, round((1-min(endeud,1))*50 + min(roa/20,1)*50, 1)))
            ieo = min(100, max(0, round(min(productiv/500000,1)*60 + (1-min(endeud,1))*40, 1)))
            idc = min(100, max(0, round(min(var_vtas/100,1)*60 + min(var_empl/50,1)*40, 1)))
            iie = min(100, max(0, round(exp_cod/4*100, 1)))
            ipt = min(100, max(0, round(min(productiv/500000,1)*50 + min(roa/20,1)*50, 1)))
            ssg = round(ice*0.25 + isf*0.20 + ieo*0.20 + idc*0.15 + iie*0.10 + ipt*0.10, 1)
            st.session_state.update({'ICE':ice,'ISF':isf,'IEO':ieo,'IDC':idc,'IIE':iie,'IPT':ipt,'SSG':ssg})
        resp = sb.table('respuestas').select('*').eq('empresa_codigo', ec).execute().data or []
        if resp:
            df_r = pd.DataFrame(resp)
            df_r['valor'] = pd.to_numeric(df_r['valor'], errors='coerce')
            df_r['bloque'] = pd.to_numeric(df_r['bloque'], errors='coerce')
            for b in range(1, 6):
                sub = df_r[df_r['bloque'] == b]['valor'].dropna()
                if len(sub) > 0:
                    st.session_state[f'score_b{b}'] = round(float(sub.mean()), 2)
    except Exception:
        pass

if not st.session_state.get('SSG'):
    _cargar_indices()

st.set_page_config(page_title="Benchmarking & Análisis", layout="wide")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Rajdhani:wght@400;600;700&display=swap');
.report-header {
    background:linear-gradient(135deg,#0a0e1a,#1a1a2e);
    border:1px solid #1e3a5f; border-radius:14px;
    padding:24px 28px; margin-bottom:8px;
}
.section-title {
    font-family:'Rajdhani',sans-serif; font-size:1.1rem; font-weight:700; color:#e2e8f0;
    background:linear-gradient(90deg,rgba(0,212,255,0.2),rgba(0,212,255,0));
    border-left:5px solid #00d4ff; padding:9px 16px; border-radius:0 8px 8px 0;
    margin:20px 0 10px 0;
}
.info-box {
    background:#0a1020; border-left:4px solid #00d4ff;
    border-radius:0 10px 10px 0; padding:12px 18px; margin-bottom:12px;
    color:#e2e8f0; line-height:1.75; font-size:.88rem;
}
.brecha-card {
    border-radius:10px; padding:11px 15px; margin-bottom:6px;
    color:#e2e8f0; line-height:1.6; font-size:.87rem;
}
.ge-card-mini {
    border-radius:10px; padding:14px 12px; text-align:center; color:#e2e8f0;
}
.sim-result {
    border-radius:12px; padding:18px 22px; text-align:center; color:#e2e8f0; margin-bottom:10px;
}
</style>
""", unsafe_allow_html=True)

INN_COLS  = ['MACRO_INNOVACION','IND_IDi','IND_GPROY','IND_DESPROD','IND_ESTRINN','IND_DESMPINN']
INN_NAMES = {'MACRO_INNOVACION':'Innovación Global','IND_IDi':'I+D+i',
             'IND_GPROY':'Gest. Proyectos','IND_DESPROD':'Dllo. Productos',
             'IND_ESTRINN':'Estrategia Inn.','IND_DESMPINN':'Desempeño Inn.'}

IDX_COLS  = ['SSG','ICE','ISF','IEO','IDC','IIE','IPT']
IDX_NAMES = {'SSG':'Score Global','ICE':'Competitividad','ISF':'Solidez Financiera',
             'IEO':'Eficiencia Operativa','IDC':'Dinamismo',
             'IIE':'Int. Exportadora','IPT':'Productividad Talento'}

ECO_COLS  = ['ROA','Var_Ventas_5a','Var_Emp_5a','Prod_Venta_Emp','Coste_Med_Emp','Ratio_Endeudamiento']
ECO_NAMES = {'ROA':'ROA (%)','Var_Ventas_5a':'Crec. Ventas 5a (%)','Var_Emp_5a':'Crec. Empleo 5a (%)',
             'Prod_Venta_Emp':'Productividad (€/emp)','Coste_Med_Emp':'Coste Medio Emp (€)',
             'Ratio_Endeudamiento':'Endeudamiento'}
ECO_SS    = {'ROA':'save_roa','Var_Ventas_5a':'save_var_vtas','Var_Emp_5a':'save_var_empl',
             'Prod_Venta_Emp':'save_productiv','Coste_Med_Emp':'save_coste_emp',
             'Ratio_Endeudamiento':'save_endeud'}
ECO_INV   = {'Ratio_Endeudamiento','Coste_Med_Emp'}

SUB_COLS = {
    'IND_IDi':     [('SUB_1.1_Dpto','Dpto. y Recursos I+D'),('SUB_1.2_PresupID','Presupuesto I+D'),('SUB_1.3_Gasto','Gasto Innovación')],
    'IND_GPROY':   [('SUB_2.1_Gbasica','Gestión Básica'),('SUB_2.2_Gavanz','Gestión Avanzada'),('SUB_2.3_OrgProy','Organización'),('SUB_2.4_EvalRdto','Evaluación Rdto.')],
    'IND_DESPROD': [('SUB_3.1_EstrDes','Estrategia Prod.'),('SUB_3.2_OportMdo','Oportunidad Mdo.'),('SUB_3.3_RecepNprod','Receptividad'),('SUB_3.4_Clientes','Orient. Cliente'),('SUB_3.5_ImpacNprod','Viabilidad')],
    'IND_ESTRINN': [('SUB_4.1_InnEstrat','Inn. Estratégica'),('SUB_4.2_CultInn','Cultura Inn.'),('SUB_4.3_Obstac','Gestión Obstác.'),('SUB_4.4_InnAb','Inn. Abierta'),('SUB_4.5_Creativ','Creatividad')],
    'IND_DESMPINN':[('SUB_5.1_ImpEstim','Impacto Estimado'),('SUB_5.2_InnEfect','Efecto Real Inn.')],
}
INN_SS = {'IND_IDi':'score_b1','IND_GPROY':'score_b2','IND_DESPROD':'score_b3',
          'IND_ESTRINN':'score_b4','IND_DESMPINN':'score_b5'}

MACROSECTORES = {1:"Industria",2:"Tec. Avanzada",3:"Servicios"}
REGIONES = {1:"Andalucía",2:"Aragón",3:"Asturias",4:"Baleares",5:"Canarias",
            6:"Cantabria",7:"C. La Mancha",8:"C. y León",9:"Cataluña",
            10:"C. Valenciana",11:"Extremadura",12:"Galicia",13:"Madrid",
            14:"Murcia",15:"Navarra",16:"País Vasco"}

PESOS_ICE = [('Prod_Venta_Emp',.30),('Var_Ventas_5a',.25),('Var_Emp_5a',.15),('Ratio_Endeudamiento',.15),('Exportacion',.15)]
PESOS_ISF = [('Ratio_Endeudamiento',.40),('ROA',.35),('Coste_Med_Emp',.25)]
PESOS_IEO = [('Prod_Venta_Emp',.50),('Coste_Med_Emp',.30),('Var_Ventas_5a',.20)]
PESOS_IDC = [('Var_Ventas_5a',.40),('Var_Emp_5a',.35),('ROA',.25)]
PESOS_IIE = [('Exportacion',.60),('Var_Ventas_5a',.20),('Prod_Venta_Emp',.20)]
PESOS_IPT = [('Prod_Venta_Emp',.40),('Coste_Med_Emp',.30),('Var_Emp_5a',.30)]
INV_ECO   = {'Ratio_Endeudamiento','Coste_Med_Emp'}

GE_LABELS = ['🚀 Líderes','⭐ Sólidas','📊 Intermedias','🔄 En Desarrollo','📉 Rezagadas']
GE_COLORS = ['#10b981','#3b82f6','#f59e0b','#f97316','#ef4444']
GE_BG     = ['#0a2010','#0a1020','#1a1200','#1a0e00','#1a0505']

LAYOUT = dict(paper_bgcolor='#0a0e1a', plot_bgcolor='#0a0e1a',
              font=dict(color='#e2e8f0', size=12), margin=dict(t=40,b=20,l=10,r=10))
hoy = date.today().strftime("%d/%m/%Y")

PERFIL_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)),'..','perfil_empresa.json')
def cargar_perfil():
    if os.path.exists(PERFIL_FILE):
        with open(PERFIL_FILE,'r',encoding='utf-8') as f: return json.load(f)
    return {}
for k,v in cargar_perfil().items():
    if k not in st.session_state: st.session_state[k] = v

if not st.session_state.get('save_reg_user'):
    st.warning("⚠️ Primero completa tu perfil en Inicio."); st.stop()

mac_cod    = int(st.session_state.get('save_macro_cod',1))
reg_cod    = int(st.session_state.get('save_reg_user',1))
nom_sector = st.session_state.get('save_sector_nombre','—')
nom_tam    = st.session_state.get('save_tam_nombre','—')
nom_reg    = st.session_state.get('save_reg_nombre','—')
nom_mac    = MACROSECTORES.get(mac_cod,'—')

def gv(k, d=0): return st.session_state.get(k,d) or d

ruta = os.path.join(os.path.dirname(os.path.abspath(__file__)),'..','datos.xlsx')
if not os.path.exists(ruta): ruta = "datos.xlsx"
try:
    df = pd.read_excel(ruta); df.columns = df.columns.str.strip()
except Exception as e:
    st.error(f"Error: {e}"); st.stop()

def calc_df(data, pesos, inv):
    s = pd.Series(0.0, index=data.index)
    for col,w in pesos:
        if col not in data.columns: continue
        s += data[col].rank(ascending=(col not in inv), pct=True)*100*w
    return s

@st.cache_data
def preparar(data):
    d = data.copy()
    d['ICE'] = calc_df(d,PESOS_ICE,INV_ECO)
    d['ISF'] = calc_df(d,PESOS_ISF,INV_ECO)
    d['IEO'] = calc_df(d,PESOS_IEO,INV_ECO)
    d['IDC'] = calc_df(d,PESOS_IDC,INV_ECO)
    d['IIE'] = calc_df(d,PESOS_IIE,INV_ECO)
    d['IPT'] = calc_df(d,PESOS_IPT,INV_ECO)
    d['SSG'] = d['ICE']*.25+d['ISF']*.20+d['IEO']*.20+d['IDC']*.15+d['IIE']*.10+d['IPT']*.10
    gs = (d['MACRO_INNOVACION'].rank(pct=True)*.25+d['ROA'].rank(pct=True)*.25+
          d['Var_Ventas_5a'].rank(pct=True)*.20+d['Prod_Venta_Emp'].rank(pct=True)*.20+
          (1-d['Ratio_Endeudamiento'].rank(pct=True))*.10)
    d['GE_score'] = gs*100
    d['GE'] = 4 - pd.qcut(gs,q=5,labels=[0,1,2,3,4]).astype(int)
    return d

df = preparar(df)

mi_inn = {
    'MACRO_INNOVACION': round((gv('score_b1')+gv('score_b2')+gv('score_b3')+gv('score_b4')+gv('score_b5'))/5,3),
    'IND_IDi':gv('score_b1'),'IND_GPROY':gv('score_b2'),'IND_DESPROD':gv('score_b3'),
    'IND_ESTRINN':gv('score_b4'),'IND_DESMPINN':gv('score_b5'),
}
mi_eco = {col: gv(ss) for col,ss in ECO_SS.items()}

def idx_emp(pesos, inv):
    s=0.0
    for col,w in pesos:
        if col not in df.columns: continue
        val=mi_eco.get(col,0)
        if val==0: continue
        p = float(np.sum(df[col]>=val)/len(df)*100) if col in inv else float(np.sum(df[col]<=val)/len(df)*100)
        s += p*w
    return round(s,1)

mi_idx = {c:idx_emp(p,INV_ECO) for c,p in zip(
    ['ICE','ISF','IEO','IDC','IIE','IPT'],
    [PESOS_ICE,PESOS_ISF,PESOS_IEO,PESOS_IDC,PESOS_IIE,PESOS_IPT])}
mi_idx['SSG'] = round(mi_idx['ICE']*.25+mi_idx['ISF']*.20+mi_idx['IEO']*.20+
                      mi_idx['IDC']*.15+mi_idx['IIE']*.10+mi_idx['IPT']*.10,1)

def ge_score_from(inn,roa,crec,prod,endeu):
    return round((float(np.sum(df['MACRO_INNOVACION']<=inn)/len(df))*.25+
                  float(np.sum(df['ROA']<=roa)/len(df))*.25+
                  float(np.sum(df['Var_Ventas_5a']<=crec)/len(df))*.20+
                  float(np.sum(df['Prod_Venta_Emp']<=prod)/len(df))*.20+
                  float(np.sum(df['Ratio_Endeudamiento']>=endeu)/len(df))*.10)*100,1)

mi_ge_score = ge_score_from(mi_inn['MACRO_INNOVACION'],mi_eco.get('ROA',0),
                             mi_eco.get('Var_Ventas_5a',0),mi_eco.get('Prod_Venta_Emp',0),
                             mi_eco.get('Ratio_Endeudamiento',1))
mi_ge = min(4,max(0,4-int(mi_ge_score/20)))

def top25(base):
    gs = (base['MACRO_INNOVACION'].rank(pct=True)*.25+base['ROA'].rank(pct=True)*.25+
          base['Var_Ventas_5a'].rank(pct=True)*.20+base['Prod_Venta_Emp'].rank(pct=True)*.20+
          (1-base['Ratio_Endeudamiento'].rank(pct=True))*.10)
    return base[gs>=gs.quantile(0.75)]

b_mac = df[df['Macrosector']==mac_cod] if len(df[df['Macrosector']==mac_cod])>=20 else df.copy()
b_reg = df[df['Region']==reg_cod]      if len(df[df['Region']==reg_cod])>=20      else df.copy()
top_mac = top25(b_mac); top_reg = top25(b_reg)

def cpct(p):
    return '#10b981' if p>=66 else '#f59e0b' if p>=33 else '#ef4444'

def pct_en(serie, val, inv=False):
    if inv: return round(float(np.sum(serie>=val)/len(serie)*100),1)
    return round(float(np.sum(serie<=val)/len(serie)*100),1)

def hex2rgb(h):
    h=h.lstrip('#')
    return int(h[0:2],16),int(h[2:4],16),int(h[4:6],16)

def radar_fig(labels, va, vb, na, nb, ca, cb, escala):
    fig = go.Figure()
    for vals,name,color,dash,op in [(va,na,ca,'solid',.20),(vb,nb,cb,'dash',.06)]:
        vc=vals+[vals[0]]; lc=labels+[labels[0]]
        r,g2,b2=hex2rgb(color); fc=f'rgba({r},{g2},{b2},{op})'
        fig.add_trace(go.Scatterpolar(r=vc,theta=lc,fill='toself',name=name,
            line=dict(color=color,width=2,dash=dash),fillcolor=fc))
    fig.update_layout(
        polar=dict(bgcolor='#0d1220',
            radialaxis=dict(range=[0,escala],gridcolor='#1e2a3a',tickfont=dict(color='#e2e8f0',size=8)),
            angularaxis=dict(gridcolor='#1e2a3a',tickfont=dict(color='#e2e8f0',size=10))),
        paper_bgcolor='#0a0e1a',font=dict(color='#e2e8f0'),
        height=360,margin=dict(t=10,b=10),
        legend=dict(bgcolor='rgba(0,0,0,0)',font=dict(color='#e2e8f0',size=10)))
    return fig

def barras_h(rows, titulo, xmax):
    fig = go.Figure()
    fig.add_trace(go.Bar(
        y=[r['n'] for r in rows], x=[r['v'] for r in rows], orientation='h',
        marker_color=[cpct(r['p']) for r in rows],
        text=[f"P{r['p']:.0f}" for r in rows],
        textposition='inside', textfont=dict(color='#fff',size=12,family='Rajdhani'),
        name='Mi Empresa'))
    fig.add_trace(go.Scatter(
        y=[r['n'] for r in rows], x=[r['t'] for r in rows],
        mode='markers', name='Media Top 25%',
        marker=dict(symbol='line-ns',size=18,color='#fff',line=dict(width=3,color='#fff'))))
    fig.update_layout(**LAYOUT,
        title=dict(text=titulo,font=dict(color='#e2e8f0',size=12)),
        height=max(240,len(rows)*52),
        xaxis=dict(range=[0,xmax],gridcolor='#1e2a3a',tickfont=dict(color='#e2e8f0')),
        yaxis=dict(gridcolor='#1e2a3a',tickfont=dict(color='#e2e8f0')),
        legend=dict(bgcolor='rgba(0,0,0,0)',font=dict(color='#e2e8f0',size=11)))
    return fig

mostrar_felix(pagina='benchmarking')
st.markdown(f"""
<div class="report-header">
  <div style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:12px;">
    <div>
      <div style="font-size:.66rem;color:#4a6fa5;letter-spacing:2px;text-transform:uppercase;margin-bottom:4px;">
        Benchmarking · Grupos Estratégicos · Mejores Prácticas · Simulador</div>
      <h1 style="font-family:Rajdhani,sans-serif;color:#00d4ff;margin:0;font-size:1.9rem;">{nom_sector}</h1>
      <p style="color:#94a3b8;margin:4px 0 0;font-size:.84rem;">{nom_tam} · {nom_reg} · {nom_mac} · {hoy}</p>
    </div>
    <div style="background:rgba(0,0,0,.4);border:1px solid {GE_COLORS[mi_ge]}60;border-radius:12px;padding:12px 20px;text-align:center;">
      <div style="font-size:.6rem;color:#64748b;letter-spacing:1px;text-transform:uppercase;">Tu Grupo Estratégico</div>
      <div style="font-family:Rajdhani,sans-serif;font-size:1.5rem;font-weight:700;color:{GE_COLORS[mi_ge]};">{GE_LABELS[mi_ge]}</div>
      <div style="font-size:.7rem;color:#64748b;">Score: {mi_ge_score:.0f}/100</div>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)

tab1,tab2,tab3,tab4 = st.tabs([
    "📊  Benchmarking",
    "🎯  Tu Grupo Estratégico",
    "🏆  Mejores Prácticas",
    "🔮  Simulador + Mapa",
])

with tab1:
    ref_lbl = st.radio("Grupo de referencia:",
                       [f"Top 25% · {nom_mac}", f"Top 25% · {nom_reg}"],
                       horizontal=True, label_visibility='collapsed', key='ref1')
    top_df  = top_mac if nom_mac in ref_lbl else top_reg
    tcol    = '#00d4ff' if nom_mac in ref_lbl else '#10b981'
    nbase   = len(b_mac) if nom_mac in ref_lbl else len(b_reg)

    st.markdown(f"""<div class="info-box" style="border-color:{tcol};">
    <strong style="color:{tcol};">{ref_lbl}</strong> · {len(top_df)} empresas de referencia
    (top 25% de {nbase}) · Percentil dentro del grupo: 🟢 ≥P66 · 🟡 P33-65 · 🔴 &lt;P33
    </div>""", unsafe_allow_html=True)

    ra,rb = st.columns(2)
    with ra:
        il=[INN_NAMES[c] for c in INN_COLS]
        iv=[mi_inn.get(c,0) for c in INN_COLS]
        tv=[round(top_df[c].mean(),2) if c in top_df.columns else 0 for c in INN_COLS]
        st.plotly_chart(radar_fig(il,iv,tv,"Mi Empresa","Top 25%",tcol,'#ffffff',5),
                        use_container_width=True)
        st.caption("Innovación (0–5)")
    with rb:
        xl=[IDX_NAMES[c] for c in IDX_COLS]
        xv=[mi_idx.get(c,0) for c in IDX_COLS]
        xt=[round(top_df[c].mean(),1) if c in top_df.columns else 0 for c in IDX_COLS]
        st.plotly_chart(radar_fig(xl,xv,xt,"Mi Empresa","Top 25%",tcol,'#ffffff',100),
                        use_container_width=True)
        st.caption("Índices Competitivos (0–100)")

    st.markdown('<div class="section-title">🔬 Innovación</div>',unsafe_allow_html=True)
    ri=[{'n':INN_NAMES[c],'v':mi_inn.get(c,0),
         't':round(top_df[c].mean(),2) if c in top_df.columns else 0,
         'p':pct_en(top_df[c],mi_inn.get(c,0)) if c in top_df.columns else 0}
        for c in INN_COLS]
    st.plotly_chart(barras_h(ri,"Innovación vs Top 25%",5.5),use_container_width=True)

    st.markdown('<div class="section-title">📊 Índices Competitivos</div>',unsafe_allow_html=True)
    rx=[{'n':IDX_NAMES[c],'v':mi_idx.get(c,0),
         't':round(top_df[c].mean(),1) if c in top_df.columns else 0,
         'p':pct_en(top_df[c],mi_idx.get(c,0)) if c in top_df.columns else 0}
        for c in IDX_COLS]
    st.plotly_chart(barras_h(rx,"Índices vs Top 25%",105),use_container_width=True)

    st.markdown('<div class="section-title">💰 Variables Económicas</div>',unsafe_allow_html=True)
    re=[]
    for c in ECO_COLS:
        if c not in top_df.columns: continue
        inv=c in ECO_INV; mv=mi_eco.get(c,0); tm=round(top_df[c].mean(),2)
        re.append({'n':ECO_NAMES[c],'v':mv,'t':tm,'p':pct_en(top_df[c],mv,inv)})
    st.plotly_chart(barras_h(re,"Económico vs Top 25%",105),use_container_width=True)
    st.caption("La barra es tu valor · El marcador blanco | es la media del top 25%")

    st.markdown('<div class="section-title">🎯 Resumen Ejecutivo</div>',unsafe_allow_html=True)
    todas=[(r['n'],r['p'],'Inn.') for r in ri]+[(r['n'],r['p'],'Comp.') for r in rx]+[(r['n'],r['p'],'Eco.') for r in re]
    forts=sorted([(n,p,t) for n,p,t in todas if p>=60],key=lambda x:-x[1])
    prios=sorted([(n,p,t) for n,p,t in todas if p<40], key=lambda x:x[1])
    cf,cp=st.columns(2)
    with cf:
        st.markdown('<p style="color:#10b981;font-weight:700;">✅ Por encima de la media del grupo</p>',unsafe_allow_html=True)
        for n,p,t in forts[:7]:
            st.markdown(f'<div class="brecha-card" style="background:#0a2010;border:1px solid #10b98140;"><span style="color:#10b981;font-weight:700;">P{p:.0f}</span> · <span style="color:#e2e8f0;">{n}</span> <span style="color:#475569;font-size:.8rem;">({t})</span></div>',unsafe_allow_html=True)
    with cp:
        st.markdown('<p style="color:#ef4444;font-weight:700;">⚠️ Áreas de mejora prioritarias</p>',unsafe_allow_html=True)
        for n,p,t in prios[:7]:
            st.markdown(f'<div class="brecha-card" style="background:#1a0505;border:1px solid #ef444440;"><span style="color:#ef4444;font-weight:700;">P{p:.0f}</span> · <span style="color:#e2e8f0;">{n}</span> <span style="color:#475569;font-size:.8rem;">({t})</span></div>',unsafe_allow_html=True)

with tab2:
    st.markdown(f"""<div style="background:{GE_BG[mi_ge]};border:2px solid {GE_COLORS[mi_ge]};
        border-radius:14px;padding:22px 28px;margin-bottom:20px;text-align:center;">
      <div style="font-size:.66rem;color:#64748b;letter-spacing:2px;text-transform:uppercase;">Tu posición en el mapa estratégico</div>
      <div style="font-family:Rajdhani,sans-serif;font-size:2.8rem;font-weight:700;color:{GE_COLORS[mi_ge]};margin:4px 0;">{GE_LABELS[mi_ge]}</div>
      <div style="color:#94a3b8;font-size:.88rem;">Score compuesto: <strong style="color:{GE_COLORS[mi_ge]};">{mi_ge_score:.0f}/100</strong>
       · Calculado sobre innovación, rentabilidad, crecimiento, productividad y endeudamiento</div>
    </div>""",unsafe_allow_html=True)

    gcols=st.columns(5)
    for i,(gc,lbl,color,bg) in enumerate(zip(gcols,GE_LABELS,GE_COLORS,GE_BG)):
        sub=df[df['GE']==i]
        activo=(i==mi_ge)
        borde=f"2px solid {color}" if activo else f"1px solid {color}35"
        sombra=f"0 0 16px {color}50" if activo else "none"
        gc.markdown(f"""<div class="ge-card-mini" style="background:{bg};border:{borde};box-shadow:{sombra};">
            <div style="font-size:.8rem;font-weight:700;color:{color};margin-bottom:5px;">{'▶ ' if activo else ''}{lbl}</div>
            <div style="font-size:.7rem;color:#64748b;margin-bottom:7px;">n={len(sub)}</div>
            <div style="font-size:.77rem;color:#e2e8f0;line-height:1.85;">
              🔬 {sub['MACRO_INNOVACION'].mean():.2f}/5<br>
              💰 ROA {sub['ROA'].mean():.1f}%<br>
              📈 {sub['Var_Ventas_5a'].mean():.1f}%<br>
              ⚙️ {sub['Prod_Venta_Emp'].mean():,.0f}€
            </div>
        </div>""",unsafe_allow_html=True)

    st.markdown("<br>",unsafe_allow_html=True)

    ge_sup=max(0,mi_ge-1)
    g_df  =df[df['GE']==mi_ge]
    s_df  =df[df['GE']==ge_sup]

    def norm100(serie,val):
        mn,mx=serie.min(),serie.max()
        return round((val-mn)/(mx-mn)*100,1) if mx>mn else 50.0

    rvars=['MACRO_INNOVACION','ROA','Var_Ventas_5a','Prod_Venta_Emp','SSG']
    rnoms=['Innovación','ROA','Crec.Ventas','Productividad','Score Global']

    rg1,rg2=st.columns(2)
    with rg1:
        if mi_ge>0:
            st.markdown(f'<p style="color:#e2e8f0;font-weight:600;margin-bottom:6px;">Perfil: <span style="color:{GE_COLORS[mi_ge]};">{GE_LABELS[mi_ge]}</span> vs <span style="color:{GE_COLORS[ge_sup]};">{GE_LABELS[ge_sup]}</span></p>',unsafe_allow_html=True)
            gv1=[norm100(df[v],g_df[v].mean()) for v in rvars]
            gv2=[norm100(df[v],s_df[v].mean()) for v in rvars]
            fig_ge=radar_fig(rnoms,gv1,gv2,GE_LABELS[mi_ge],GE_LABELS[ge_sup],GE_COLORS[mi_ge],GE_COLORS[ge_sup],100)
            st.plotly_chart(fig_ge,use_container_width=True)
        else:
            st.success("🏆 Eres Líder. No hay grupo superior al que aspirar — consolida tu ventaja.")

    with rg2:
        if mi_ge>0:
            st.markdown(f'<p style="color:#e2e8f0;font-weight:600;margin-bottom:8px;">¿Qué necesitas para ascender a <span style="color:{GE_COLORS[ge_sup]};">{GE_LABELS[ge_sup]}</span>?</p>',unsafe_allow_html=True)
            ascenso=[
                ('MACRO_INNOVACION','Innovación (0-5)',mi_inn['MACRO_INNOVACION'],False),
                ('ROA','ROA (%)',mi_eco.get('ROA',0),False),
                ('Var_Ventas_5a','Crec. Ventas 5a (%)',mi_eco.get('Var_Ventas_5a',0),False),
                ('Prod_Venta_Emp','Productividad (€/emp)',mi_eco.get('Prod_Venta_Emp',0),False),
                ('Ratio_Endeudamiento','Endeudamiento',mi_eco.get('Ratio_Endeudamiento',1),True),
            ]
            for col,nom,mival,inv in ascenso:
                obj=s_df[col].quantile(0.25) if not inv else s_df[col].quantile(0.75)
                ok=(mival>=obj) if not inv else (mival<=obj)
                brecha=abs(obj-mival)
                cc='#10b981' if ok else '#f59e0b' if brecha/(abs(obj)+.001)<0.3 else '#ef4444'
                st.markdown(f"""<div class="brecha-card" style="background:#0d1b2a;border-left:3px solid {cc};border:1px solid {cc}30;border-left:3px solid {cc};">
                    <div style="display:flex;justify-content:space-between;flex-wrap:wrap;gap:4px;">
                      <span style="color:#e2e8f0;font-weight:600;">{'✅' if ok else '⚠️'} {nom}</span>
                      <span style="color:{cc};font-size:.85rem;font-weight:700;">
                        {'Ya lo superas' if ok else f'Objetivo: {obj:.2f}'}
                      </span>
                    </div>
                    <div style="color:#64748b;font-size:.78rem;margin-top:3px;">
                      Tu valor: <strong style="color:#e2e8f0;">{mival:.2f}</strong>
                      · Umbral grupo superior: <strong style="color:{GE_COLORS[ge_sup]};">{obj:.2f}</strong>
                      {f'· Brecha: {brecha:.2f}' if not ok else ' · ✓ Cumplido'}
                    </div>
                </div>""",unsafe_allow_html=True)
        else:
            st.info("Mantén tu liderazgo reforzando las dimensiones con menor margen sobre la media.")

with tab3:
    st.markdown("""<div class="info-box">
    Para cada bloque de innovación, compara tus subindicadores con el <strong style="color:#00d4ff;">top 25%</strong>
    y con la media global. Los subindicadores con mayor diferencial entre el top y la media son las
    <strong style="color:#10b981;">prácticas diferenciales</strong> — ahí se concentra la ventaja competitiva real.
    </div>""",unsafe_allow_html=True)

    col_sel,col_ref=st.columns([2,2])
    with col_sel:
        blq=st.selectbox("Bloque de innovación:",list(INN_SS.keys()),
                         format_func=lambda x:INN_NAMES[x],key='blq')
    with col_ref:
        ref_mp=st.radio("Referencia:",
                        [f"Top 25% · {nom_mac}",f"Top 25% · {nom_reg}"],
                        horizontal=True,key='ref_mp')
    top_mp=top_mac if nom_mac in ref_mp else top_reg

    subs=SUB_COLS.get(blq,[])
    if subs:
        scols=[s[0] for s in subs]; snoms=[s[1] for s in subs]
        mi_blq=gv(INN_SS[blq])
        top_m =[round(top_mp[c].mean(),3) if c in top_mp.columns else 0 for c in scols]
        glob_m=[round(df[c].mean(),3)     if c in df.columns else 0     for c in scols]
        mi_sv =[gv(f'score_sub{blq[-1]}_{i+1}', glob_m[i]) for i in range(len(subs))]

        fig_mp=go.Figure()
        for vals,name,color in [(mi_sv,"Mi Empresa",'#00d4ff'),(top_m,"Top 25%",'#10b981'),(glob_m,"Media Global",'#475569')]:
            fig_mp.add_trace(go.Bar(y=snoms,x=vals,orientation='h',name=name,
                marker_color=color,opacity=0.85,
                text=[f"{v:.2f}" for v in vals],textposition='outside',
                textfont=dict(color='#e2e8f0',size=10)))
        fig_mp.update_layout(**LAYOUT,barmode='group',
            title=dict(text=f"Subindicadores · {INN_NAMES[blq]}",font=dict(color='#e2e8f0',size=12)),
            height=max(300,len(subs)*85),
            xaxis=dict(range=[0,5.8],gridcolor='#1e2a3a',tickfont=dict(color='#e2e8f0')),
            yaxis=dict(gridcolor='#1e2a3a',tickfont=dict(color='#e2e8f0')),
            legend=dict(bgcolor='rgba(0,0,0,0)',font=dict(color='#e2e8f0',size=11)))
        st.plotly_chart(fig_mp,use_container_width=True)

        st.markdown('<div class="section-title">💡 Prácticas Diferenciales</div>',unsafe_allow_html=True)
        difs=sorted([(snoms[i],top_m[i]-glob_m[i],top_m[i]-mi_sv[i],top_m[i],glob_m[i],mi_sv[i])
                     for i in range(len(subs))],key=lambda x:-x[1])
        for nom,dif,mib,tv,gv2,miv in difs:
            ok=mib<=0; cc='#10b981' if ok else '#f59e0b' if mib<0.4 else '#ef4444'
            barra_pct=min(100,miv/5*100)
            st.markdown(f"""<div class="brecha-card" style="background:#0d1b2a;border-left:4px solid {cc};border:1px solid {cc}25;border-left:4px solid {cc};">
                <div style="display:flex;justify-content:space-between;flex-wrap:wrap;gap:4px;">
                  <span style="color:#e2e8f0;font-weight:600;">{'✅' if ok else '⚠️'} {nom}</span>
                  <span style="color:#64748b;font-size:.82rem;">
                    Diferencial top·media: <strong style="color:#00d4ff;">+{dif:.2f}</strong>
                    &nbsp;·&nbsp; Tu brecha vs top: <strong style="color:{cc};">{mib:+.2f}</strong>
                  </span>
                </div>
                <div style="background:#1e2a3a;border-radius:4px;height:5px;width:100%;margin-top:8px;">
                  <div style="background:{cc};border-radius:4px;height:5px;width:{barra_pct:.0f}%;"></div>
                </div>
                <div style="display:flex;justify-content:space-between;font-size:.75rem;color:#64748b;margin-top:3px;">
                  <span>Tu valor: <strong style="color:#e2e8f0;">{miv:.2f}</strong></span>
                  <span>Media global: {gv2:.2f}</span>
                  <span style="color:#10b981;">Top 25%: {tv:.2f}</span>
                </div>
            </div>""",unsafe_allow_html=True)

with tab4:
    sc,mc=st.columns([1,1],gap="large")

    with sc:
        st.markdown('<div class="section-title">🔮 Simulador de Escenarios</div>',unsafe_allow_html=True)
        st.markdown('<p style="color:#94a3b8;font-size:.84rem;">Mueve los sliders para ver cómo cambia tu score y grupo estratégico en tiempo real.</p>',unsafe_allow_html=True)

        with st.expander("🔬 Indicadores de Innovación", expanded=True):
            sb1=st.slider("I+D+i",              0.0,5.0,float(mi_inn.get('IND_IDi',2.5)),0.1,key='sb1')
            sb2=st.slider("Gestión Proyectos",  0.0,5.0,float(mi_inn.get('IND_GPROY',2.5)),0.1,key='sb2')
            sb3=st.slider("Dllo. Productos",    0.0,5.0,float(mi_inn.get('IND_DESPROD',2.5)),0.1,key='sb3')
            sb4=st.slider("Estrategia Inn.",    0.0,5.0,float(mi_inn.get('IND_ESTRINN',2.5)),0.1,key='sb4')
            sb5=st.slider("Desempeño Inn.",     0.0,5.0,float(mi_inn.get('IND_DESMPINN',2.5)),0.1,key='sb5')
        sim_inn_val=round((sb1+sb2+sb3+sb4+sb5)/5,3)

        with st.expander("💰 Variables Económicas", expanded=True):
            def slrange(col,default):
                mn=float(df[col].quantile(.05)); mx=float(df[col].quantile(.95))
                dv=float(mi_eco.get(col,default)); dv=max(mn,min(mx,dv))
                step=round((mx-mn)/100,2) or 0.01
                return mn,mx,dv,step
            mn,mx,dv,st_=slrange('ROA',8);        sim_roa  =st.slider("ROA (%)",mn,mx,dv,st_,key='sroa')
            mn,mx,dv,st_=slrange('Var_Ventas_5a',30); sim_crec =st.slider("Crec. Ventas 5a (%)",mn,mx,dv,st_,key='screc')
            mn,mx,dv,st_=slrange('Var_Emp_5a',20);    sim_cemp =st.slider("Crec. Empleo 5a (%)",mn,mx,dv,st_,key='scemp')
            mn,mx,dv,st_=slrange('Prod_Venta_Emp',200000); sim_prod=st.slider("Productividad (€)",mn,mx,dv,max(st_,1000.0),key='sprod')
            mn,mx,dv,st_=slrange('Coste_Med_Emp',40000);   sim_coemp=st.slider("Coste Medio Emp (€)",mn,mx,dv,max(st_,500.0),key='scoemp')
            mn,mx,dv,st_=slrange('Ratio_Endeudamiento',0.6); sim_end=st.slider("Endeudamiento",mn,mx,dv,max(st_,0.01),key='send')

        sim_score=ge_score_from(sim_inn_val,sim_roa,sim_crec,sim_prod,sim_end)
        sim_ge=min(4,max(0,4-int(sim_score/20)))
        delta=round(sim_score-mi_ge_score,1)
        delta_ge=mi_ge-sim_ge
        csim=GE_COLORS[sim_ge]
        dtxt=f"▲ +{delta}" if delta>0 else f"▼ {delta}" if delta<0 else "= Sin cambio"
        dcol='#10b981' if delta>0 else '#ef4444' if delta<0 else '#64748b'
        ascenso_txt=f'<div style="margin-top:8px;background:rgba(16,185,129,.15);border-radius:8px;padding:6px 12px;color:#10b981;font-size:.88rem;font-weight:600;">🎉 Ascenso: {GE_LABELS[mi_ge]} → {GE_LABELS[sim_ge]}</div>' if delta_ge>0 else ''
        descenso_txt=f'<div style="margin-top:8px;background:rgba(239,68,68,.12);border-radius:8px;padding:6px 12px;color:#ef4444;font-size:.88rem;">⬇️ Descenso: {GE_LABELS[mi_ge]} → {GE_LABELS[sim_ge]}</div>' if delta_ge<0 else ''

        st.markdown(f"""<div class="sim-result" style="background:{GE_BG[sim_ge]};border:2px solid {csim};">
            <div style="font-size:.62rem;color:#64748b;letter-spacing:1.5px;text-transform:uppercase;">Resultado del Escenario</div>
            <div style="font-family:Rajdhani,sans-serif;font-size:2.3rem;font-weight:700;color:{csim};margin:4px 0;">{GE_LABELS[sim_ge]}</div>
            <div style="font-size:1rem;color:#e2e8f0;">Score: <strong style="color:{csim};">{sim_score:.0f}/100</strong>
              &nbsp;·&nbsp; <span style="color:{dcol};font-weight:700;">{dtxt} pts</span></div>
            {ascenso_txt}{descenso_txt}
        </div>""",unsafe_allow_html=True)

        fig_vel=go.Figure(go.Indicator(
            mode="gauge+number+delta",
            value=sim_score,
            delta={'reference':mi_ge_score,'valueformat':'.1f',
                   'increasing':{'color':'#10b981'},'decreasing':{'color':'#ef4444'}},
            number={'font':{'size':30,'color':csim},'valueformat':'.1f'},
            gauge={'axis':{'range':[0,100],'tickcolor':'#334155','tickfont':{'color':'#e2e8f0','size':9}},
                   'bar':{'color':csim,'thickness':0.28},'bgcolor':'#0a0e1a','borderwidth':0,
                   'steps':[{'range':[0,20],'color':'#1a0505'},{'range':[20,40],'color':'#1a0e00'},
                             {'range':[40,60],'color':'#1a1200'},{'range':[60,80],'color':'#0a1020'},
                             {'range':[80,100],'color':'#0a2010'}],
                   'threshold':{'line':{'color':'#fff','width':2},'thickness':0.75,'value':mi_ge_score}}))
        fig_vel.update_layout(height=200,margin=dict(l=20,r=20,t=10,b=10),paper_bgcolor='rgba(0,0,0,0)')
        st.plotly_chart(fig_vel,use_container_width=True)

    with mc:
        st.markdown('<div class="section-title">🗺️ Mapa de Posicionamiento Sectorial</div>',unsafe_allow_html=True)
        st.markdown('<p style="color:#94a3b8;font-size:.84rem;">X = Innovación · Y = Score Competitivo · ★ = Tu empresa · ✦ = Escenario simulado</p>',unsafe_allow_html=True)

        dmap=df[df['Macrosector']==mac_cod].copy()
        med_inn=dmap['MACRO_INNOVACION'].mean(); med_ssg=dmap['GE_score'].mean()

        fig_map=go.Figure()
        for gn in range(5):
            sub=dmap[dmap['GE']==gn]
            if len(sub)==0: continue
            fig_map.add_trace(go.Scatter(
                x=sub['MACRO_INNOVACION'],y=sub['GE_score'],mode='markers',name=GE_LABELS[gn],
                marker=dict(color=GE_COLORS[gn],size=5,opacity=0.5,line=dict(width=0.2,color='#000')),
                hovertemplate=f"{GE_LABELS[gn]}<br>Inn: %{{x:.2f}}<br>Score: %{{y:.1f}}<extra></extra>"))

        fig_map.add_trace(go.Scatter(
            x=[mi_inn['MACRO_INNOVACION']],y=[mi_ge_score],mode='markers+text',
            name='Tú (actual)',showlegend=True,
            marker=dict(symbol='star',size=22,color='#ffffff',line=dict(width=2,color=GE_COLORS[mi_ge])),
            text=['★ Tú'],textposition='top right',textfont=dict(color='#ffffff',size=11)))

        if abs(sim_score-mi_ge_score)>1 or abs(sim_inn_val-mi_inn['MACRO_INNOVACION'])>0.05:
            fig_map.add_trace(go.Scatter(
                x=[sim_inn_val],y=[sim_score],mode='markers+text',
                name='Tú (simulado)',showlegend=True,
                marker=dict(symbol='star',size=22,color=csim,line=dict(width=2,color='#ffffff')),
                text=['✦ Sim.'],textposition='top right',textfont=dict(color=csim,size=11)))
            fig_map.add_annotation(
                x=sim_inn_val,y=sim_score,ax=mi_inn['MACRO_INNOVACION'],ay=mi_ge_score,
                xref='x',yref='y',axref='x',ayref='y',
                arrowhead=2,arrowsize=1.2,arrowwidth=2,arrowcolor=csim)

        cuad_data=[
            (0,med_inn,med_ssg,100,'Competitivos\nno innovadores','#3b82f6'),
            (med_inn,5.2,med_ssg,100,'Líderes\nEstratégicos','#10b981'),
            (0,med_inn,0,med_ssg,'Rezagados','#ef4444'),
            (med_inn,5.2,0,med_ssg,'Innovadores\nno rentables','#f59e0b'),
        ]
        for x0,x1,y0,y1,txt,c in cuad_data:
            r2,g3,b3=hex2rgb(c)
            fig_map.add_shape(type='rect',x0=x0,x1=x1,y0=y0,y1=y1,
                line=dict(color=c,width=1,dash='dot'),
                fillcolor=f'rgba({r2},{g3},{b3},0.04)')
            fig_map.add_annotation(x=(x0+x1)/2,y=(y0+y1)/2,text=txt,
                showarrow=False,font=dict(color=c,size=9,family='Rajdhani'),opacity=0.55)

        fig_map.update_layout(**LAYOUT,height=580,
            xaxis=dict(title='Índice de Innovación (0-5)',range=[0,5.3],
                       gridcolor='#1e2a3a',tickfont=dict(color='#e2e8f0'),
                       title_font=dict(color='#e2e8f0')),
            yaxis=dict(title='Score Competitivo (0-100)',range=[0,105],
                       gridcolor='#1e2a3a',tickfont=dict(color='#e2e8f0'),
                       title_font=dict(color='#e2e8f0')),
            legend=dict(bgcolor='rgba(0,0,0,0)',font=dict(color='#e2e8f0',size=10),
                        orientation='h',yanchor='bottom',y=1.01,xanchor='left',x=0))
        st.plotly_chart(fig_map,use_container_width=True)

        inn_alto=mi_inn['MACRO_INNOVACION']>=med_inn; ssg_alto=mi_ge_score>=med_ssg
        ctxt={
            (True,True): ("Líderes Estratégicos","#10b981","Alta innovación y alto desempeño competitivo. Posición de referencia en el sector."),
            (True,False):("Innovadores no Rentables","#f59e0b","Alta innovación pero desempeño económico bajo la media. La inversión en innovación aún no se ha traducido en resultados."),
            (False,True):("Competitivos no Innovadores","#3b82f6","Buen desempeño pero baja innovación. Posición vulnerable a medio-largo plazo."),
            (False,False):("Rezagados","#ef4444","Por debajo de la media en ambas dimensiones. Mayor margen de mejora en el sector."),
        }
        cn,cc2,cdesc=ctxt[(inn_alto,ssg_alto)]
        st.markdown(f"""<div class="info-box" style="border-color:{cc2};">
            <strong style="color:{cc2};">Tu cuadrante actual: {cn}</strong><br>
            <span style="color:#e2e8f0;">{cdesc}</span>
        </div>""",unsafe_allow_html=True)

# ── DESCARGAS ─────────────────────────────────────────────────────────────
st.divider()
st.markdown('<div class="section-title">📥 Descargar Informe de Benchmarking</div>', unsafe_allow_html=True)

from datetime import date as _date
import io as _io

def generar_html_benchmarking():
    hoy_dl = _date.today().strftime("%d/%m/%Y")
    sector_dl = st.session_state.get('save_sector_nombre','—')
    tam_dl    = st.session_state.get('save_tam_nombre','—')
    reg_dl    = st.session_state.get('save_reg_nombre','—')
    b1_dl = round(st.session_state.get('score_b1',0),2)
    b2_dl = round(st.session_state.get('score_b2',0),2)
    b3_dl = round(st.session_state.get('score_b3',0),2)
    b4_dl = round(st.session_state.get('score_b4',0),2)
    b5_dl = round(st.session_state.get('score_b5',0),2)
    ssg_dl = round(st.session_state.get('SSG',0),1)
    ice_dl = round(st.session_state.get('ICE',0),1)
    isf_dl = round(st.session_state.get('ISF',0),1)
    ieo_dl = round(st.session_state.get('IEO',0),1)
    idc_dl = round(st.session_state.get('IDC',0),1)
    iie_dl = round(st.session_state.get('IIE',0),1)
    ipt_dl = round(st.session_state.get('IPT',0),1)
    nivel = lambda v, thr1=70, thr2=40: 'Alto' if v > thr1 else ('Medio' if v > thr2 else 'Bajo')
    nivel5 = lambda v: 'Alto' if v > 3.5 else ('Medio' if v > 2.5 else 'Bajo')
    ssg_nivel = 'Posicion solida' if ssg_dl > 70 else ('Posicion intermedia' if ssg_dl > 50 else 'Posicion debil')
    macro = round((b1_dl+b2_dl+b3_dl+b4_dl+b5_dl)/5, 2) if any([b1_dl,b2_dl,b3_dl,b4_dl,b5_dl]) else 0
    if macro > 2.75 and ssg_dl > 50:
        cuadrante = 'Lideres Estrategicos'
    elif macro > 2.75:
        cuadrante = 'Innovadores no Rentables'
    elif ssg_dl > 50:
        cuadrante = 'Competitivos no Innovadores'
    else:
        cuadrante = 'Rezagados'
    css = "<style>body{font-family:Georgia,serif;max-width:920px;margin:50px auto;color:#1a1a1a;line-height:1.75;}h1{color:#1d4ed8;border-bottom:3px solid #1d4ed8;padding-bottom:10px;}h2{color:#1d4ed8;border-left:4px solid #1d4ed8;padding-left:12px;margin-top:28px;}.perfil{background:#f0f7ff;border-radius:8px;padding:14px 18px;margin:14px 0;}table{border-collapse:collapse;width:100%;margin:14px 0;font-size:.9rem;}th{background:#1d4ed8;color:white;padding:10px;text-align:left;}td{padding:9px 11px;border-bottom:1px solid #e5e7eb;}.cuadrante{background:#eff6ff;border-left:4px solid #1d4ed8;padding:12px 16px;margin:8px 0;border-radius:0 8px 8px 0;}.nota{border:1px dashed #9ca3af;border-radius:8px;padding:36px 16px;margin:20px 0;color:#9ca3af;font-style:italic;text-align:center;}</style>"
    indices_rows = ""
    for nombre, val in [('SSG Score Global', str(ssg_dl)+'/100'),('ICE Competitividad', str(ice_dl)+'/100'),
        ('ISF Solidez Financiera', str(isf_dl)+'/100'),('IEO Eficiencia Operativa', str(ieo_dl)+'/100'),
        ('IDC Dinamismo', str(idc_dl)+'/100'),('IIE Exportacion', str(iie_dl)+'/100'),('IPT Productividad', str(ipt_dl)+'/100')]:
        n_val = float(val.split('/')[0])
        indices_rows += "<tr><td>" + nombre + "</td><td><strong>" + val + "</strong></td><td>" + nivel(n_val) + "</td></tr>"
    inn_rows = ""
    for nombre, val in [('I+D+i', b1_dl),('Gestion de Proyectos', b2_dl),
        ('Desarrollo de Productos', b3_dl),('Estrategia de Innovacion', b4_dl),('Desempeno de Innovacion', b5_dl)]:
        inn_rows += "<tr><td>" + nombre + "</td><td>" + str(val) + "/5</td><td>" + nivel5(val) + "</td></tr>"
    html = "<!DOCTYPE html><html><head><meta charset='utf-8'>" + css + "</head><body>"
    html += "<h1>Informe de Benchmarking Estrategico</h1>"
    html += "<div class='perfil'><strong>Empresa:</strong> " + sector_dl + " · " + tam_dl + " · " + reg_dl + " · <strong>Fecha:</strong> " + hoy_dl + "</div>"
    html += "<h2>1. Posicion Competitiva Global</h2><p>" + ssg_nivel + " (SSG: " + str(ssg_dl) + "/100)</p>"
    html += "<table><tr><th>Indice</th><th>Valor</th><th>Nivel</th></tr>" + indices_rows + "</table>"
    html += "<h2>2. Perfil de Innovacion</h2>"
    html += "<table><tr><th>Indicador</th><th>Puntuacion</th><th>Nivel</th></tr>" + inn_rows + "</table>"
    html += "<h2>3. Posicion Estrategica</h2><div class='cuadrante'><strong>Cuadrante:</strong> " + cuadrante + "</div>"
    html += "<h2>4. Notas del equipo directivo</h2><div class='nota'>Espacio para anotaciones</div>"
    html += "<hr/><p style='color:#9ca3af;font-size:.78rem;text-align:center;'>Etelvia · " + hoy_dl + "</p></body></html>"
    return html

def generar_word_benchmarking():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        hoy_dl = _date.today().strftime("%d/%m/%Y")
        sector_dl = st.session_state.get('save_sector_nombre','—')
        tam_dl    = st.session_state.get('save_tam_nombre','—')
        reg_dl    = st.session_state.get('save_reg_nombre','—')
        b1_dl = round(st.session_state.get('score_b1',0),2)
        b2_dl = round(st.session_state.get('score_b2',0),2)
        b3_dl = round(st.session_state.get('score_b3',0),2)
        b4_dl = round(st.session_state.get('score_b4',0),2)
        b5_dl = round(st.session_state.get('score_b5',0),2)
        ssg_dl = round(st.session_state.get('SSG',0),1)
        ice_dl = round(st.session_state.get('ICE',0),1)
        isf_dl = round(st.session_state.get('ISF',0),1)
        ieo_dl = round(st.session_state.get('IEO',0),1)
        idc_dl = round(st.session_state.get('IDC',0),1)
        iie_dl = round(st.session_state.get('IIE',0),1)
        ipt_dl = round(st.session_state.get('IPT',0),1)
        nivel = lambda v: 'Alto' if v > 70 else ('Medio' if v > 40 else 'Bajo')
        nivel5 = lambda v: 'Alto' if v > 3.5 else ('Medio' if v > 2.5 else 'Bajo')
        doc = Document()
        t = doc.add_heading('Informe de Benchmarking Estrategico', 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Empresa: " + sector_dl + " · " + tam_dl + " · " + reg_dl + " · " + hoy_dl)
        doc.add_heading('1. Posicion Competitiva Global', level=1)
        tabla = doc.add_table(rows=1, cols=3)
        tabla.style = 'Table Grid'
        hdr = tabla.rows[0].cells
        hdr[0].text = 'Indice'; hdr[1].text = 'Valor'; hdr[2].text = 'Nivel'
        for nombre, val in [('SSG Score Global', ssg_dl),('ICE Competitividad', ice_dl),
            ('ISF Solidez Financiera', isf_dl),('IEO Eficiencia Operativa', ieo_dl),
            ('IDC Dinamismo', idc_dl),('IIE Exportacion', iie_dl),('IPT Productividad', ipt_dl)]:
            row = tabla.add_row().cells
            row[0].text = nombre; row[1].text = str(val)+'/100'; row[2].text = nivel(val)
        doc.add_heading('2. Perfil de Innovacion', level=1)
        tabla2 = doc.add_table(rows=1, cols=3)
        tabla2.style = 'Table Grid'
        hdr2 = tabla2.rows[0].cells
        hdr2[0].text = 'Indicador'; hdr2[1].text = 'Puntuacion'; hdr2[2].text = 'Nivel'
        for nombre, val in [('I+D+i', b1_dl),('Gestion Proyectos', b2_dl),
            ('Desarrollo Productos', b3_dl),('Estrategia Innovacion', b4_dl),('Desempeno', b5_dl)]:
            row = tabla2.add_row().cells
            row[0].text = nombre; row[1].text = str(val)+'/5'; row[2].text = nivel5(val)
        doc.add_heading('3. Notas y comentarios', level=1)
        for _ in range(8):
            doc.add_paragraph('_' * 80)
        buf = _io.BytesIO()
        doc.save(buf); buf.seek(0)
        return buf.getvalue()
    except ImportError:
        return None

html_bench = generar_html_benchmarking()
word_bench = generar_word_benchmarking()

col1, col2, col3 = st.columns(3)
with col1:
    st.download_button("📄 Descargar Word (.docx)", data=word_bench if word_bench else b"",
        file_name="benchmarking_estrategico.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary", use_container_width=True, disabled=word_bench is None)
with col2:
    st.download_button("🌐 Descargar HTML", data=html_bench,
        file_name="benchmarking_estrategico.html", mime="text/html", use_container_width=True)
with col3:
    st.info("Para PDF: abre el HTML → **Ctrl+P** → **Guardar como PDF**")