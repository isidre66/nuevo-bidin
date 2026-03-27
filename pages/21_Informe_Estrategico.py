import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import json, os, numpy as np, requests, unicodedata
from streamlit_echarts import st_echarts
from datetime import date

def cargar_api_key():
    try:
        return st.secrets["ANTHROPIC_API_KEY"]
    except Exception: pass
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    if key: return key
    env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '.env')
    if os.path.exists(env_path):
        with open(env_path, 'r') as f:
            for line in f:
                if line.startswith("ANTHROPIC_API_KEY"):
                    return line.split("=", 1)[1].strip().strip('"').strip("'")
    return ""

ANTHROPIC_API_KEY = cargar_api_key()

st.set_page_config(page_title="Informe Estratégico-Competitivo", layout="wide")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Rajdhani:wght@400;600;700&family=Inter:wght@300;400;500&display=swap');
.report-header {
    background:linear-gradient(135deg,#0a0e1a,#1a2744);
    border:1px solid #1e3a5f; border-radius:14px;
    padding:28px 32px; margin-bottom:24px;
}
.section-title {
    font-family:'Rajdhani',sans-serif; font-size:1.2rem; font-weight:700; color:#fff;
    background:linear-gradient(90deg,rgba(0,212,255,0.2),rgba(0,212,255,0));
    border-left:5px solid #00d4ff; padding:10px 16px; border-radius:0 8px 8px 0;
    margin:28px 0 14px 0; text-shadow:0 0 12px rgba(0,212,255,0.6);
}
.text-block {
    background:linear-gradient(135deg,#0d1b2a,#1a2744);
    border:1px solid #1e3a5f; border-radius:12px;
    padding:22px 26px; margin-bottom:16px;
    color:#cbd5e1; line-height:1.8; font-size:.93rem;
}
.rec-block {
    background:linear-gradient(135deg,#1a1200,#2a1e00);
    border-left:4px solid #f59e0b; border-radius:0 10px 10px 0;
    padding:18px 22px; margin-bottom:14px;
    color:#cbd5e1; line-height:1.8; font-size:.93rem;
}
.pred-block {
    background:linear-gradient(135deg,#0a0a2e,#1a1035);
    border-left:4px solid #7c3aed; border-radius:0 10px 10px 0;
    padding:18px 22px; margin-bottom:14px;
    color:#cbd5e1; line-height:1.8; font-size:.93rem;
}
.pending {
    color:#475569; font-style:italic; background:#0d1b2a;
    border:1px dashed #1e3a5f; border-radius:10px;
    padding:18px 22px; margin-bottom:14px;
}
</style>
""", unsafe_allow_html=True)

# ── Perfil local ──────────────────────────────────────────────────────────
PERFIL_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)),'..','perfil_empresa.json')
def cargar_perfil():
    if os.path.exists(PERFIL_FILE):
        with open(PERFIL_FILE,'r',encoding='utf-8') as f: return json.load(f)
    return {}
perfil = cargar_perfil()
for k,v in perfil.items():
    if k not in st.session_state: st.session_state[k] = v

# ── Cargar desde Supabase si falta perfil ────────────────────────────────
if not st.session_state.get('save_reg_user'):
    try:
        from supabase import create_client
        _url = st.secrets.get("SUPABASE_URL", os.environ.get("SUPABASE_URL",""))
        _key = st.secrets.get("SUPABASE_KEY", os.environ.get("SUPABASE_KEY",""))
        _ec  = st.session_state.get('empresa_codigo','') or perfil.get('sesion_codigo','')
        if _url and _key and _ec:
            _sb = create_client(_url, _key)
            _ed = _sb.table('empresas').select('*').eq('codigo',_ec).execute().data
            if _ed:
                _e = _ed[0]
                _SM = {"Alimentación y bebidas":1,"Textil y confección":2,"Cuero y calzado":3,
                    "Química y plásticos":4,"Minerales no metálicos":5,"Metalmecanico":6,
                    "Maquinaria equipo":7,"Otras manufacturas":8,"Electrónica, telecomunicaciones":9,
                    "Informática, software. Robótica, IA":10,"Actividades I+D: biotech, farmacia":11,
                    "Transporte y logística":12,"Consultoría y servicios profesionales":13,
                    "Turismo y hosteleria":14,"Retail y comercio":15,"Otros servicios":16}
                _MM = {1:1,2:1,3:1,4:1,5:1,6:1,7:1,8:1,9:2,10:2,11:2,12:3,13:3,14:3,15:3,16:3}
                _sc = _SM.get(_e.get('sector',''),0)
                st.session_state.update({
                    'save_sector_nombre':_e.get('sector',''),'save_tam_nombre':_e.get('tamano',''),
                    'save_reg_nombre':_e.get('region',''),'save_export_nombre':_e.get('exportacion',''),
                    'save_anti_nombre':_e.get('antiguedad',''),
                    'save_ventas':float(_e.get('ventas') or 0),
                    'save_empleados':int(_e.get('empleados') or 0),
                    'save_roa':float(_e.get('roa') or 0),
                    'save_var_vtas':float(_e.get('var_ventas') or 0),
                    'save_var_empl':float(_e.get('var_empleados') or 0),
                    'save_productiv':float(_e.get('productividad') or 0),
                    'save_coste_emp':float(_e.get('coste_empleado') or 0),
                    'save_endeud':float(_e.get('endeudamiento') or 0),
                    'save_reg_user':{"Andalucia":1,"Aragon":2,"Asturias":3,"Baleares":4,"Canarias":5,
                        "Cantabria":6,"Castilla la Mancha":7,"Castilla y León":8,"Cataluña":9,
                        "Com Valenciana":10,"Extremadura":11,"Galicia":12,"Madrid":13,
                        "Murcia":14,"Navarra":15,"Pais Vasco":16}.get(_e.get('region',''),0),
                    'save_tam_user':{"Pequeña":1,"Mediana":2,"Grande":3}.get(_e.get('tamano',''),0),
                    'save_sector_cod':_sc,'save_macro_cod':_MM.get(_sc,3),
                    'save_export_cod':{"Menos 10 %":1,"10 - 30 %":2,"30 - 60 %":3,"> 60 %":4}.get(_e.get('exportacion',''),0),
                    'save_anti_cod':{"Menos 10 años":1,"10-30 años":2,"> 30 años":3}.get(_e.get('antiguedad',''),0),
                    'empresa_codigo': _ec,
                })
    except Exception: pass

if not st.session_state.get('save_reg_user'):
    st.warning("⚠️ Primero completa tu perfil en **Mi Empresa → Datos de la empresa**.")
    st.stop()

# ── Excel ─────────────────────────────────────────────────────────────────
ruta = os.path.join(os.path.dirname(os.path.abspath(__file__)),'..','datos.xlsx')
if not os.path.exists(ruta): ruta = "datos.xlsx"
try:
    df = pd.read_excel(ruta); df.columns = df.columns.str.strip()
except Exception as e:
    st.error(f"Error: {e}"); st.stop()

EXPORTACIONES = {1:"Menos 10%",2:"10-30%",3:"30-60%",4:">60%"}
nom_sector = st.session_state.get('save_sector_nombre','—')
nom_reg    = st.session_state.get('save_reg_nombre','—')
nom_tam    = st.session_state.get('save_tam_nombre','—')
nom_export = EXPORTACIONES.get(st.session_state.get('save_export_cod',1),'—')
nom_anti   = st.session_state.get('save_anti_nombre','—')
mac_cod    = st.session_state.get('save_macro_cod',1)
tam_cod    = st.session_state.get('save_tam_user',1)
hoy        = date.today().strftime("%d/%m/%Y")

dff = df[(df['Macrosector']==mac_cod)&(df['Tamaño']==tam_cod)]
if len(dff) < 10: dff = df[df['Macrosector']==mac_cod]
if len(dff) < 10: dff = df.copy()
n_ref = len(dff)

mis = {
    'Ventas':              st.session_state.get('save_ventas',0),
    'Empleados':           st.session_state.get('save_empleados',0),
    'ROA':                 st.session_state.get('save_roa',0),
    'Var_Ventas_5a':       st.session_state.get('save_var_vtas',0),
    'Var_Emp_5a':          st.session_state.get('save_var_empl',0),
    'Prod_Venta_Emp':      st.session_state.get('save_productiv',0),
    'Coste_Med_Emp':       st.session_state.get('save_coste_emp',0),
    'Ratio_Endeudamiento': st.session_state.get('save_endeud',0),
    'Exportacion':         st.session_state.get('save_export_cod',1),
}

def pct(series, val, inv=False):
    if inv: return float(np.sum(series >= val)/len(series)*100)
    return float(np.sum(series <= val)/len(series)*100)

def col_v(v):
    if v >= 66: return '#10b981'
    if v >= 33: return '#f59e0b'
    return '#ef4444'

def nivel(v):
    if v >= 66: return 'alto'
    if v >= 33: return 'medio'
    return 'bajo'

def norm(s):
    return ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn').upper()

def calc_idx(df_r, d):
    def p(col, inv=False):
        if col not in df_r.columns: return 50.0
        return pct(df_r[col], d.get(col, df_r[col].mean()), inv)
    ICE = p('Prod_Venta_Emp')*.30+p('Var_Ventas_5a')*.25+p('Var_Emp_5a')*.15+p('Ratio_Endeudamiento',True)*.15+p('Exportacion')*.15
    ISF = p('Ratio_Endeudamiento',True)*.40+p('ROA')*.35+p('Coste_Med_Emp',True)*.25
    IEO = p('Prod_Venta_Emp')*.50+p('Coste_Med_Emp',True)*.30+p('Var_Ventas_5a')*.20
    IDC = p('Var_Ventas_5a')*.40+p('Var_Emp_5a')*.35+p('ROA')*.25
    IIE = p('Exportacion')*.60+p('Var_Ventas_5a')*.20+p('Prod_Venta_Emp')*.20
    IPT = p('Prod_Venta_Emp')*.40+p('Coste_Med_Emp')*.30+p('Var_Emp_5a')*.30
    SSG = ICE*.25+ISF*.20+IEO*.20+IDC*.15+IIE*.10+IPT*.10
    return {k:round(v,1) for k,v in {'ICE':ICE,'ISF':ISF,'IEO':IEO,'IDC':IDC,'IIE':IIE,'IPT':IPT,'SSG':SSG}.items()}

idx = calc_idx(dff, mis)

pcts_eco = {
    'Prod_Venta_Emp':      round(pct(dff['Prod_Venta_Emp'], mis['Prod_Venta_Emp'])),
    'ROA':                 round(pct(dff['ROA'], mis['ROA'])),
    'Ratio_Endeudamiento': round(pct(dff['Ratio_Endeudamiento'], mis['Ratio_Endeudamiento'], inv=True)),
    'Var_Ventas_5a':       round(pct(dff['Var_Ventas_5a'], mis['Var_Ventas_5a'])),
    'Var_Emp_5a':          round(pct(dff['Var_Emp_5a'], mis['Var_Emp_5a'])),
}
grp = {c: round(dff[c].mean(),2) for c in ['Prod_Venta_Emp','ROA','Ratio_Endeudamiento','Var_Ventas_5a','Var_Emp_5a','Coste_Med_Emp']}

# ── IA ────────────────────────────────────────────────────────────────────
def llamar_ia():
    if not ANTHROPIC_API_KEY:
        return "ERROR: No se encontró la API key de Anthropic."
    prompt = f"""Eres un consultor estratégico senior. Redacta un informe ejecutivo de diagnóstico competitivo en español para esta empresa. Máximo 700 palabras totales. Escribe en párrafos narrativos fluidos, sin viñetas ni listas en el texto.

PERFIL: {nom_sector} | {nom_tam} | {nom_reg} | Exportación: {nom_export} | Antigüedad: {nom_anti}
Grupo de referencia: {n_ref} empresas del mismo macrosector y tamaño.

ÍNDICES ESTRATÉGICOS (0-100, percentil vs grupo):
SSG={idx['SSG']} ({nivel(idx['SSG'])}) | ICE={idx['ICE']} | ISF={idx['ISF']} | IEO={idx['IEO']} | IDC={idx['IDC']} | IIE={idx['IIE']} | IPT={idx['IPT']}

VARIABLES ECONÓMICAS (empresa vs media grupo):
- Productividad: {mis['Prod_Venta_Emp']:,.0f}€/emp vs {grp['Prod_Venta_Emp']:,.0f}€/emp (P{pcts_eco['Prod_Venta_Emp']})
- ROA: {mis['ROA']:.1f}% vs {grp['ROA']:.1f}% (P{pcts_eco['ROA']})
- Endeudamiento: {mis['Ratio_Endeudamiento']:.2f} vs {grp['Ratio_Endeudamiento']:.2f} (P{pcts_eco['Ratio_Endeudamiento']} invertido)
- Crecimiento ventas 5a: {mis['Var_Ventas_5a']:.1f}% vs {grp['Var_Ventas_5a']:.1f}% (P{pcts_eco['Var_Ventas_5a']})
- Crecimiento empleo 5a: {mis['Var_Emp_5a']:.1f}% vs {grp['Var_Emp_5a']:.1f}% (P{pcts_eco['Var_Emp_5a']})

TONO Y ESTILO — OBLIGATORIO:
- Escribe como un consultor estratégico senior (estilo McKinsey, BCG). Tono sobrio, preciso y constructivo.
- PROHIBIDO: "alarmante", "catastrófico", "grave", "urgente", "preocupante", "peligroso", "severo", "excelente", "fantástico", "brillante", "impresionante".
- USA EN SU LUGAR: "por debajo de la media del sector", "margen de mejora relevante", "área de desarrollo prioritario", "posición sólida", "ventaja competitiva sostenida", "brecha que merece atención", "oportunidad de mejora", "inferior/superior a la media del grupo".
- Las debilidades se presentan como oportunidades de mejora, no como fracasos.

ESTRUCTURA OBLIGATORIA — usa exactamente estos 4 títulos con ##:

## POSICION COMPETITIVA GLOBAL
2 párrafos. Qué posición ocupa la empresa, qué perfil competitivo tiene. Menciona SSG y 2-3 índices clave con sus valores.

## ANALISIS POR DIMENSIONES
3 párrafos: Párrafo 1 — Eficiencia y Competitividad (ICE+IEO). Párrafo 2 — Solidez Financiera (ISF). Párrafo 3 — Dinamismo y Apertura exterior (IDC+IIE).

## FORTALEZAS Y RIESGOS
2 párrafos. Primer párrafo: máximo 3 fortalezas con datos. Segundo párrafo: máximo 3 áreas de mejora.

## RECOMENDACIONES Y PREDICCION
3 recomendaciones numeradas con horizonte temporal. Luego un párrafo de predicción si actúa vs si no actúa en 2-3 años.

IMPORTANTE: Escribe los títulos SIN tildes tal como aparecen arriba."""
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"Content-Type":"application/json","x-api-key":ANTHROPIC_API_KEY,"anthropic-version":"2023-06-01"},
            json={"model":"claude-sonnet-4-20250514","max_tokens":1500,"messages":[{"role":"user","content":prompt}]},
            timeout=45
        )
        data = resp.json()
        if 'content' in data and data['content']: return data['content'][0]['text']
        return f"ERROR: {data.get('error',{}).get('message','Respuesta inesperada')}"
    except requests.exceptions.Timeout: return "ERROR: Tiempo de espera agotado."
    except Exception as e: return f"ERROR: {e}"

def parsear(texto):
    resultado = []
    partes = [p.strip() for p in texto.split('##') if p.strip()]
    for parte in partes:
        lineas = parte.split('\n', 1)
        titulo_orig = lineas[0].strip()
        cuerpo = lineas[1].strip() if len(lineas) > 1 else ''
        resultado.append((norm(titulo_orig), titulo_orig, cuerpo))
    return resultado

def buscar_seccion(secciones, *palabras_clave):
    for titulo_norm, titulo_orig, cuerpo in secciones:
        for kw in palabras_clave:
            if norm(kw) in titulo_norm: return cuerpo
    return None

def render(texto, estilo='normal'):
    if not texto: return ''
    html = texto.replace('\n\n','</p><p>').replace('\n','<br>')
    css_class = {'normal':'text-block','rec':'rec-block','pred':'pred-block'}.get(estilo,'text-block')
    return f'<div class="{css_class}"><p>{html}</p></div>'

# ── Cabecera ──────────────────────────────────────────────────────────────
ssg_c = col_v(idx['SSG'])
st.markdown(f"""
<div class="report-header">
  <div style="display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:16px;">
    <div>
      <div style="font-size:.7rem;color:#64748b;letter-spacing:2px;text-transform:uppercase;margin-bottom:4px;">
        Informe Estratégico · Diagnóstico Competitivo</div>
      <h1 style="font-family:Rajdhani,sans-serif;color:#00d4ff;margin:0;font-size:2rem;">{nom_sector}</h1>
      <p style="color:#94a3b8;margin:6px 0 0 0;font-size:.88rem;">
        {nom_tam} · {nom_reg} · Exportación: {nom_export} · {nom_anti}<br>
        <span style="color:#64748b;font-size:.8rem;">Referencia: {n_ref} empresas comparables · {hoy}</span>
      </p>
    </div>
    <div style="text-align:center;background:rgba(0,0,0,0.35);border:1px solid {ssg_c}55;border-radius:12px;padding:16px 28px;">
      <div style="font-size:.65rem;color:#64748b;letter-spacing:1.5px;text-transform:uppercase;">Score Estratégico Global</div>
      <div style="font-family:Rajdhani,sans-serif;font-size:4.2rem;font-weight:700;color:{ssg_c};line-height:1;">{idx['SSG']}</div>
      <div style="font-size:.78rem;color:#94a3b8;">/ 100 · {nivel(idx['SSG']).upper()}</div>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)

# ── Velocímetros ──────────────────────────────────────────────────────────
st.markdown('<div class="section-title">⚡ Los 6 Índices Estratégicos</div>', unsafe_allow_html=True)
gauge_defs = [('ICE','Competitividad'),('ISF','Solidez Fin.'),('IEO','Eficiencia Op.'),
              ('IDC','Dinamismo'),('IIE','Internacionalización'),('IPT','Productividad')]
gcols = st.columns(6)
for col_st,(cod,nombre) in zip(gcols, gauge_defs):
    v = idx[cod]
    fig = go.Figure(go.Indicator(
        mode="gauge+number", value=v,
        number={'font':{'size':26,'color':col_v(v)},'valueformat':'.1f'},
        title={'text':f"{cod}<br><span style='font-size:9px;color:#64748b;'>{nombre}</span>",'font':{'size':11,'color':'#94a3b8'}},
        gauge={
            'axis':{'range':[0,100],'tickcolor':'#334155','tickfont':{'color':'#64748b','size':7}},
            'bar':{'color':col_v(v),'thickness':0.28},'bgcolor':'#0a0e1a','borderwidth':0,
            'steps':[{'range':[0,33],'color':'#1a0505'},{'range':[33,66],'color':'#1a1200'},{'range':[66,100],'color':'#051a0a'}],
            'threshold':{'line':{'color':'#475569','width':2},'thickness':0.75,'value':50}
        }
    ))
    fig.update_layout(height=190,margin=dict(l=8,r=8,t=42,b=4),paper_bgcolor='rgba(0,0,0,0)')
    col_st.plotly_chart(fig, use_container_width=True)

st.divider()
col_btn, col_hint = st.columns([1,2])
with col_btn:
    generar = st.button("🚀 Generar Informe Completo con IA", type="primary", use_container_width=True)
with col_hint:
    st.info("La IA redacta ~700 palabras analizando tus índices y variables económicas reales. Tiempo estimado: 15-25 segundos.")

if generar:
    with st.spinner("Redactando informe personalizado..."):
        st.session_state['inf1_texto'] = llamar_ia()

texto_raw = st.session_state.get('inf1_texto', '')
secciones = parsear(texto_raw) if texto_raw else []
if texto_raw.startswith('ERROR:'):
    st.error(f"❌ {texto_raw}"); st.stop()

# ── Secciones ─────────────────────────────────────────────────────────────
st.markdown('<div class="section-title">🎯 Posición Competitiva Global</div>', unsafe_allow_html=True)
txt = buscar_seccion(secciones, 'POSICION', 'GLOBAL', 'COMPETITIVA')
st.markdown(render(txt) if txt else '<div class="pending">Pulsa "Generar Informe" para ver el análisis.</div>', unsafe_allow_html=True)

st.markdown('<div class="section-title">📊 Comparativa Visual vs. Grupo</div>', unsafe_allow_html=True)
c_radar, c_rank = st.columns(2)
idx_labs = ['ICE','ISF','IEO','IDC','IIE','IPT']
emp_r = [idx[k] for k in idx_labs]

with c_radar:
    radar_opt = {
        "backgroundColor":"#0a0e1a",
        "tooltip":{"trigger":"item","backgroundColor":"#1e293b","textStyle":{"color":"#e2e8f0"}},
        "legend":{"data":["Mi Empresa","Mediana Grupo (50)"],"textStyle":{"color":"#94a3b8"},"bottom":0},
        "radar":{"indicator":[{"name":k,"max":100} for k in idx_labs],"shape":"polygon","splitNumber":4,
            "axisName":{"color":"#94a3b8","fontSize":12},
            "splitLine":{"lineStyle":{"color":["#1e3a5f"]*4}},
            "splitArea":{"areaStyle":{"color":["rgba(0,212,255,0.02)","rgba(0,212,255,0.05)"]}},
            "axisLine":{"lineStyle":{"color":"#1e3a5f"}}},
        "series":[{"type":"radar","data":[
            {"value":emp_r,"name":"Mi Empresa","itemStyle":{"color":"#00d4ff"},
             "lineStyle":{"color":"#00d4ff","width":2},"areaStyle":{"color":"rgba(0,212,255,0.18)"},
             "symbol":"circle","symbolSize":6},
            {"value":[50]*6,"name":"Mediana Grupo (50)","itemStyle":{"color":"#f59e0b"},
             "lineStyle":{"color":"#f59e0b","width":1.5,"type":"dashed"},
             "areaStyle":{"color":"rgba(245,158,11,0.05)"}},
        ]}],
    }
    st_echarts(options=radar_opt, height="360px")

with c_rank:
    pct_names = ['ICE · Competitividad','ISF · Solidez Fin.','IEO · Eficiencia Op.',
                 'IDC · Dinamismo','IIE · Internacionalización','IPT · Productividad']
    rank_opt = {
        "backgroundColor":"#0a0e1a",
        "tooltip":{"trigger":"axis","axisPointer":{"type":"shadow"},"backgroundColor":"#1e293b","textStyle":{"color":"#e2e8f0"}},
        "grid":{"left":"2%","right":"14%","bottom":"5%","containLabel":True},
        "xAxis":{"type":"value","min":0,"max":100,"axisLabel":{"color":"#64748b"},"splitLine":{"lineStyle":{"color":"#1e3a5f"}}},
        "yAxis":{"type":"category","data":pct_names,"inverse":True,"axisLabel":{"color":"#94a3b8","fontSize":11}},
        "series":[{"type":"bar",
            "data":[{"value":v,"itemStyle":{"color":col_v(v),"borderRadius":[0,5,5,0]}} for v in emp_r],
            "label":{"show":True,"position":"right","color":"#e2e8f0","fontSize":13,"fontFamily":"Rajdhani","formatter":"{c}"},
            "markLine":{"silent":True,"lineStyle":{"color":"#f59e0b","type":"dashed","width":1.5},
                "data":[{"xAxis":50,"label":{"formatter":"Mediana","color":"#f59e0b","fontSize":9}}]},
        }]
    }
    st_echarts(options=rank_opt, height="360px")

st.markdown('<div class="section-title">📐 Análisis por Dimensiones</div>', unsafe_allow_html=True)
txt = buscar_seccion(secciones, 'ANALISIS', 'DIMENSION')
st.markdown(render(txt) if txt else '<div class="pending">Análisis por dimensiones pendiente.</div>', unsafe_allow_html=True)

st.markdown('<div class="section-title">⚖️ Fortalezas y Riesgos</div>', unsafe_allow_html=True)
cf, cg4 = st.columns([1,1])
with cf:
    txt = buscar_seccion(secciones, 'FORTALEZA', 'RIESGO', 'DEBILIDAD')
    st.markdown(render(txt) if txt else '<div class="pending">Fortalezas y riesgos pendientes.</div>', unsafe_allow_html=True)
with cg4:
    st.markdown("**📍 Tu posición: Productividad vs. ROA**")
    prod_grp = dff['Prod_Venta_Emp'].values.tolist() if 'Prod_Venta_Emp' in dff.columns else []
    roa_grp  = dff['ROA'].values.tolist() if 'ROA' in dff.columns else []
    scatter_data = [[round(float(x),1), round(float(y),1)] for x,y in zip(prod_grp, roa_grp)]
    scatter_opt = {
        "backgroundColor":"#0a0e1a",
        "tooltip":{"trigger":"item","backgroundColor":"#1e293b","textStyle":{"color":"#e2e8f0"},
            "formatter":"function(p){ return p.seriesName+'<br/>Prod: '+p.value[0].toLocaleString()+'€<br/>ROA: '+p.value[1]+'%'; }"},
        "legend":{"data":[f"Grupo ({n_ref})","Mi Empresa"],"textStyle":{"color":"#94a3b8"},"top":5},
        "xAxis":{"name":"Productividad (€/emp)","nameTextStyle":{"color":"#64748b"},"type":"value",
            "axisLabel":{"color":"#64748b","formatter":"function(v){return (v/1000).toFixed(0)+'k';}"},
            "splitLine":{"lineStyle":{"color":"#1e3a5f"}}},
        "yAxis":{"name":"ROA (%)","nameTextStyle":{"color":"#64748b"},"type":"value",
            "axisLabel":{"color":"#64748b"},"splitLine":{"lineStyle":{"color":"#1e3a5f"}}},
        "dataZoom":[{"type":"inside"}],
        "series":[
            {"name":f"Grupo ({n_ref})","type":"scatter","data":scatter_data,"symbolSize":5,
             "itemStyle":{"color":"rgba(16,185,129,0.4)","borderColor":"rgba(16,185,129,0.6)","borderWidth":1}},
            {"name":"Mi Empresa","type":"scatter","data":[[round(float(mis['Prod_Venta_Emp']),1),round(float(mis['ROA']),1)]],
             "symbolSize":24,"symbol":"star",
             "itemStyle":{"color":"#00d4ff","borderColor":"white","borderWidth":2,"shadowBlur":20,"shadowColor":"rgba(0,212,255,0.8)"},
             "label":{"show":True,"formatter":"Yo","color":"#00d4ff","position":"top","fontSize":11,"fontWeight":"bold"}},
        ],
    }
    st_echarts(options=scatter_opt, height="360px")

st.markdown('<div class="section-title">🎯 Recomendaciones y Predicción</div>', unsafe_allow_html=True)
txt_rec = buscar_seccion(secciones, 'RECOMENDACION', 'PREDICCION')
if txt_rec:
    parrafos = [p.strip() for p in txt_rec.split('\n\n') if p.strip()]
    recs  = [p for p in parrafos if p and p[0].isdigit() and '.' in p[:3]]
    preds = [p for p in parrafos if p not in recs]
    if not recs and not preds:
        st.markdown(render(txt_rec, 'rec'), unsafe_allow_html=True)
    else:
        if recs: st.markdown(f'<div class="rec-block"><p>{"</p><p>".join(recs)}</p></div>', unsafe_allow_html=True)
        if preds: st.markdown(f'<div class="pred-block"><p>{"</p><p>".join(preds)}</p></div>', unsafe_allow_html=True)
else:
    c1,c2 = st.columns(2)
    c1.markdown('<div class="pending">Recomendaciones pendientes.</div>', unsafe_allow_html=True)
    c2.markdown('<div class="pending">Predicción pendiente.</div>', unsafe_allow_html=True)

st.divider()

# ── Descarga ──────────────────────────────────────────────────────────────
st.markdown('<div class="section-title">📥 Descargar Informe</div>', unsafe_allow_html=True)

def generar_html():
    ia_html = ""
    colores_sec = {'POSICION':'#0066cc','ANALISIS':'#0066cc','FORTALEZA':'#dc2626',
                   'RIESGO':'#dc2626','RECOMENDACION':'#d97706','PREDICCION':'#7c3aed'}
    for tn, to, cuerpo in secciones:
        color = '#0066cc'
        for kw,c in colores_sec.items():
            if kw in tn: color=c; break
        ia_html += f"<h2 style='color:{color};border-left:4px solid {color};padding-left:10px;margin-top:28px;'>{to}</h2><p style='line-height:1.75;'>{cuerpo.replace(chr(10),'<br>')}</p>"
    idx_tabla = "".join([f"<tr><td><strong>{cod}</strong></td><td>{desc}</td><td style='text-align:center;font-weight:700;color:{'green' if idx[cod]>=66 else 'orange' if idx[cod]>=33 else 'red'};'>{idx[cod]}/100</td></tr>"
        for cod,desc in [('SSG','Score Estratégico Global'),('ICE','Competitividad Empresarial'),
            ('ISF','Solidez Financiera'),('IEO','Eficiencia Operativa'),('IDC','Dinamismo y Crecimiento'),
            ('IIE','Intensidad Exportadora'),('IPT','Productividad y Talento')]])
    ssg_color = 'green' if idx['SSG']>=66 else 'orange' if idx['SSG']>=33 else 'red'
    return f"""<!DOCTYPE html><html><head><meta charset='utf-8'>
<style>body{{font-family:Georgia,serif;margin:50px auto;max-width:860px;color:#1a1a1a;line-height:1.75;}}
h1{{color:#0066cc;border-bottom:3px solid #0066cc;padding-bottom:10px;font-size:1.8rem;}}
.perfil{{background:#f0f7ff;border-radius:8px;padding:16px;margin:16px 0;font-size:.9rem;}}
.ssg{{font-size:2.8rem;font-weight:700;color:{ssg_color};}}
table{{border-collapse:collapse;width:100%;margin:16px 0;font-size:.9rem;}}
th{{background:#0066cc;color:white;padding:10px;text-align:left;}}
td{{padding:9px 11px;border-bottom:1px solid #e5e7eb;}}
.notas{{border:1px dashed #9ca3af;border-radius:8px;padding:40px 16px;margin:16px 0;
    color:#9ca3af;font-style:italic;font-size:.85rem;text-align:center;}}
@media print{{body{{margin:20px;}}}}</style></head><body>
<h1>Informe Estratégico · Diagnóstico Competitivo</h1>
<div class="perfil"><strong>{nom_sector}</strong> | {nom_tam} | {nom_reg} | Exportación: {nom_export} | {nom_anti}<br>
Referencia: <strong>{n_ref} empresas comparables</strong> | Fecha: {hoy}</div>
<p>Score Estratégico Global (SSG): <span class="ssg">{idx['SSG']}</span> / 100 · {nivel(idx['SSG']).upper()}</p>
<h2 style='color:#0066cc;'>Índices Estratégicos</h2>
<table><tr><th>Índice</th><th>Descripción</th><th>Valor (0-100)</th></tr>{idx_tabla}</table>
{ia_html or '<p style="color:#6b7280;font-style:italic;">Genera el análisis IA antes de descargar.</p>'}
<div class="notas">✏️ Espacio para tus anotaciones y comentarios del equipo directivo</div>
<hr style='margin-top:40px;'/>
<p style='color:#9ca3af;font-size:.78rem;text-align:center;'>Plataforma de Diagnóstico Estratégico Empresarial · {hoy}</p>
</body></html>"""

def generar_word():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        doc = Document()
        t = doc.add_heading('Informe Estratégico · Diagnóstico Competitivo', 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Empresa: {nom_sector} | {nom_tam} | {nom_reg} | Exportación: {nom_export} | {nom_anti}")
        doc.add_paragraph(f"Referencia: {n_ref} empresas comparables | Fecha: {hoy}")
        doc.add_paragraph(f"Score Estratégico Global (SSG): {idx['SSG']} / 100 — {nivel(idx['SSG']).upper()}")
        doc.add_heading('Índices Estratégicos', level=1)
        tabla = doc.add_table(rows=1, cols=3); tabla.style = 'Table Grid'
        hdr = tabla.rows[0].cells
        hdr[0].text='Índice'; hdr[1].text='Descripción'; hdr[2].text='Valor (0-100)'
        for cod,desc in [('SSG','Score Estratégico Global'),('ICE','Competitividad Empresarial'),
            ('ISF','Solidez Financiera'),('IEO','Eficiencia Operativa'),('IDC','Dinamismo y Crecimiento'),
            ('IIE','Intensidad Exportadora'),('IPT','Productividad y Talento')]:
            row=tabla.add_row().cells; row[0].text=cod; row[1].text=desc; row[2].text=f"{idx[cod]}/100"
        if secciones:
            for tn,to,cuerpo in secciones:
                doc.add_heading(to, level=1); doc.add_paragraph(cuerpo)
        else:
            doc.add_paragraph("Genera el análisis IA en la plataforma antes de descargar.")
        doc.add_heading('Notas y comentarios del equipo directivo', level=1)
        doc.add_paragraph("[ Escribe aquí tus conclusiones, observaciones y próximos pasos ]")
        for _ in range(8):
            p=doc.add_paragraph(); p.add_run('_'*80)
        buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()
    except ImportError: return None

html_out = generar_html()
word_out = generar_word()

st.markdown("""<div style="background:#f0fdf4;border:1px solid #10b981;border-radius:10px;
    padding:14px 18px;margin-bottom:16px;font-size:.87rem;color:#065f46;">
    📋 <strong>El documento Word es completamente editable</strong> — puedes añadir texto, 
    modificar el contenido e insertar el logo de tu empresa antes de enviarlo al equipo directivo.
</div>""", unsafe_allow_html=True)

c1,c2,c3 = st.columns(3)
with c1:
    st.download_button("📄 Descargar Word (.docx)", data=word_out if word_out else b"",
        file_name=f"informe_estrategico_{nom_sector[:15].replace(' ','_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary", use_container_width=True, disabled=word_out is None)
with c2:
    st.download_button("🌐 Descargar HTML", data=html_out,
        file_name=f"informe_estrategico_{nom_sector[:15].replace(' ','_')}.html",
        mime="text/html", use_container_width=True)
with c3:
    st.info("Para PDF: abre el HTML en navegador → **Ctrl+P** → **Guardar como PDF**")