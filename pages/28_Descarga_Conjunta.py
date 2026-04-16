import streamlit as st
import os
import io
from datetime import date

st.set_page_config(page_title="Descarga Conjunta de Informes", layout="wide")

if not st.session_state.get('save_reg_user'):
    st.warning("⚠️ Por favor, complete primero el perfil de su empresa en la sección Mi Empresa.")
    st.stop()

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Rajdhani:wght@400;600;700&family=Inter:wght@300;400;500&display=swap');
.dl-header {
    background: linear-gradient(135deg, #0a0e1a, #1a2744);
    border: 1px solid #1e3a5f;
    border-radius: 14px;
    padding: 28px 32px;
    margin-bottom: 24px;
}
</style>
""", unsafe_allow_html=True)

hoy = date.today().strftime("%d/%m/%Y")
sector = st.session_state.get('save_sector_nombre', '—')
tam    = st.session_state.get('save_tam_nombre', '—')
reg    = st.session_state.get('save_reg_nombre', '—')
b1 = round(st.session_state.get('score_b1', 0), 2)
b2 = round(st.session_state.get('score_b2', 0), 2)
b3 = round(st.session_state.get('score_b3', 0), 2)
b4 = round(st.session_state.get('score_b4', 0), 2)
b5 = round(st.session_state.get('score_b5', 0), 2)
macro_inn = round((b1+b2+b3+b4+b5)/5, 2) if any([b1,b2,b3,b4,b5]) else 0
ssg = round(st.session_state.get('SSG', 0), 1)
ice = round(st.session_state.get('ICE', 0), 1)
isf = round(st.session_state.get('ISF', 0), 1)
ieo = round(st.session_state.get('IEO', 0), 1)
idc = round(st.session_state.get('IDC', 0), 1)
iie = round(st.session_state.get('IIE', 0), 1)
ipt = round(st.session_state.get('IPT', 0), 1)
roa       = st.session_state.get('save_roa', 0)
var_vtas  = st.session_state.get('save_var_vtas', 0)
var_empl  = st.session_state.get('save_var_empl', 0)
productiv = st.session_state.get('save_productiv', 0)
endeud    = st.session_state.get('save_endeud', 0)
ventas    = st.session_state.get('save_ventas', 0)
empleados = st.session_state.get('save_empleados', 0)

nivel  = lambda v, t1=70, t2=40: 'Alto' if v>t1 else ('Medio' if v>t2 else 'Bajo')
nivel5 = lambda v: 'Alto' if v>3.5 else ('Medio' if v>2.5 else 'Bajo')
nivel_ssg = 'Posicion solida' if ssg>70 else ('Posicion intermedia' if ssg>50 else 'Posicion debil')
nivel_inn = 'Avanzado' if macro_inn>3.5 else ('Medio' if macro_inn>2.5 else 'Por debajo de la media')

st.markdown(f"""
<div class="dl-header">
  <div style="font-size:.7rem;color:#4a6fa5;letter-spacing:2px;text-transform:uppercase;margin-bottom:6px;">Descarga Conjunta de Informes</div>
  <h1 style="font-family:'Rajdhani',sans-serif;color:#00d4ff;margin:0;font-size:1.8rem;">Paquete completo de informes</h1>
  <p style="color:#94a3b8;margin:6px 0 0;font-size:.85rem;">{sector} · {tam} · {reg} · {hoy} · SSG: {ssg}/100 · Innovacion: {macro_inn}/5</p>
</div>
""", unsafe_allow_html=True)

st.markdown("### Selecciona los informes a incluir")
inc_estrategico  = st.checkbox("📄 Informe Estratégico Competitivo", value=True)
inc_innovacion   = st.checkbox("🔬 Informe de Innovación", value=True)
inc_benchmarking = st.checkbox("🎯 Benchmarking Estratégico", value=True)
inc_plan         = st.checkbox("🚀 Plan de Acción y Hoja de Ruta", value=True)
st.markdown("---")

def get_acciones():
    acc = []
    if b1<2.5: acc.append(("Q1","Incrementar inversion en I+D+i","Alto"))
    if b2<2.5: acc.append(("Q1","Implementar metodologias agiles","Alto"))
    if b3<2.5: acc.append(("Q2","Fortalecer desarrollo de nuevos productos","Alto"))
    if b4<2.5: acc.append(("Q2","Definir estrategia de innovacion","Medio"))
    if b5<2.5: acc.append(("Q3","Medir impacto real de la innovacion","Medio"))
    if ssg<50: acc.append(("Q1","Revisar modelo de negocio y posicionamiento","Alto"))
    if idc<40: acc.append(("Q2","Activar palancas de crecimiento en ventas","Alto"))
    if not acc:
        acc = [("Q1","Consolidar posicion competitiva","Medio"),
               ("Q2","Explorar nuevos mercados","Medio"),
               ("Q3","Reforzar capacidades de innovacion","Medio"),
               ("Q4","Establecer alianzas estrategicas","Bajo")]
    return acc[:5]

def get_cuadrante():
    macro = round((b1+b2+b3+b4+b5)/5,2) if any([b1,b2,b3,b4,b5]) else 0
    if macro>2.75 and ssg>50: return 'Lideres Estrategicos'
    elif macro>2.75: return 'Innovadores no Rentables'
    elif ssg>50: return 'Competitivos no Innovadores'
    return 'Rezagados'

def generar_html_conjunto():
    CSS = "<style>body{font-family:Georgia,serif;max-width:960px;margin:40px auto;color:#1a1a1a;line-height:1.75;}h1{color:#1d4ed8;border-bottom:3px solid #1d4ed8;padding-bottom:10px;margin-top:40px;}h2{color:#1d4ed8;border-left:4px solid #1d4ed8;padding-left:12px;margin-top:28px;}h3{color:#334155;margin-top:20px;}.perfil{background:#f0f7ff;border-radius:8px;padding:14px 18px;margin:14px 0;font-size:.9rem;}.seccion{background:#f8fafc;border:1px solid #e2e8f0;border-radius:10px;padding:20px 24px;margin:16px 0;}table{border-collapse:collapse;width:100%;margin:14px 0;font-size:.9rem;}th{background:#1d4ed8;color:white;padding:10px;text-align:left;}td{padding:9px 11px;border-bottom:1px solid #e5e7eb;}.accion{background:#eff6ff;border-left:4px solid #1d4ed8;padding:12px 16px;margin:8px 0;border-radius:0 8px 8px 0;}.nota{border:1px dashed #9ca3af;border-radius:8px;padding:36px 16px;margin:20px 0;color:#9ca3af;font-style:italic;text-align:center;}.page-break{page-break-before:always;border-top:2px solid #e2e8f0;margin-top:40px;padding-top:20px;}@media print{.page-break{page-break-before:always;}}</style>"
    html = "<!DOCTYPE html><html><head><meta charset='utf-8'>" + CSS + "</head><body>"
    html += "<h1 style='font-size:2rem;text-align:center;border:none;margin-top:60px;'>Informe Estrategico Completo</h1>"
    html += "<div class='perfil' style='text-align:center;'><strong>" + sector + " · " + tam + " · " + reg + "</strong><br>Fecha: " + hoy + " · SSG: " + str(ssg) + "/100 · Innovacion: " + str(macro_inn) + "/5</div>"

    if inc_estrategico:
        html += "<div class='page-break'><h1>1. Informe Estrategico Competitivo</h1>"
        html += "<div class='perfil'>" + sector + " · " + tam + " · " + reg + " · " + hoy + "</div>"
        html += "<h2>Indices Competitivos</h2><table><tr><th>Indice</th><th>Valor</th><th>Nivel</th></tr>"
        for n,v in [('SSG',ssg),('ICE',ice),('ISF',isf),('IEO',ieo),('IDC',idc),('IIE',iie),('IPT',ipt)]:
            html += "<tr><td>" + n + "</td><td><strong>" + str(v) + "/100</strong></td><td>" + nivel(v) + "</td></tr>"
        html += "</table><h2>Variables Economicas</h2><table><tr><th>Variable</th><th>Valor</th></tr>"
        for n,v,u in [('Ventas',ventas,'k EUR'),('Empleados',empleados,''),('ROA',roa,'%'),('Var.Ventas 5a',var_vtas,'%'),('Var.Empleo 5a',var_empl,'%'),('Productividad',productiv,'EUR/emp'),('Endeudamiento',endeud,'')]:
            html += "<tr><td>" + n + "</td><td>" + str(round(v,2)) + " " + u + "</td></tr>"
        html += "</table><div class='nota'>Espacio para anotaciones del equipo directivo</div></div>"

    if inc_innovacion:
        html += "<div class='page-break'><h1>2. Informe de Innovacion</h1>"
        html += "<h2>Indice Global: " + str(macro_inn) + "/5 — " + nivel_inn + "</h2>"
        html += "<table><tr><th>Indicador</th><th>Puntuacion</th><th>Nivel</th></tr>"
        for n,v in [('I+D+i',b1),('Gestion Proyectos',b2),('Desarrollo Productos',b3),('Estrategia Innovacion',b4),('Desempeno Innovacion',b5)]:
            html += "<tr><td>" + n + "</td><td>" + str(v) + "/5</td><td>" + nivel5(v) + "</td></tr>"
        html += "</table><div class='nota'>Espacio para anotaciones del equipo directivo</div></div>"

    if inc_benchmarking:
        html += "<div class='page-break'><h1>3. Benchmarking Estrategico</h1>"
        html += "<h2>Cuadrante: " + get_cuadrante() + " · SSG: " + str(ssg) + "/100</h2>"
        html += "<table><tr><th>Indice</th><th>Valor</th><th>Nivel</th></tr>"
        for n,v in [('SSG',ssg),('ICE',ice),('ISF',isf),('IEO',ieo),('IDC',idc),('IIE',iie),('IPT',ipt)]:
            html += "<tr><td>" + n + "</td><td>" + str(v) + "/100</td><td>" + nivel(v) + "</td></tr>"
        html += "</table><div class='nota'>Espacio para anotaciones del equipo directivo</div></div>"

    if inc_plan:
        html += "<div class='page-break'><h1>4. Plan de Accion y Hoja de Ruta</h1>"
        html += "<p>SSG: <strong>" + str(ssg) + "/100</strong> · Innovacion: <strong>" + str(macro_inn) + "/5</strong></p>"
        html += "<h2>Acciones Estrategicas Prioritarias</h2>"
        for trim, accion, impacto in get_acciones():
            html += "<div class='accion'><strong>" + trim + ":</strong> " + accion + " · Impacto: " + impacto + "</div>"
        html += "<h2>Hoja de Ruta 3 Anos</h2>"
        html += "<div class='accion'><strong>Ano 1:</strong> Consolidacion — Focalizar en brechas criticas.</div>"
        html += "<div class='accion'><strong>Ano 2:</strong> Aceleracion — Escalar iniciativas de mayor impacto.</div>"
        html += "<div class='accion'><strong>Ano 3:</strong> Liderazgo — Consolidar posicion de referencia.</div>"
        html += "<div class='nota'>Espacio para anotaciones del equipo directivo</div></div>"

    html += "<hr/><p style='color:#9ca3af;font-size:.78rem;text-align:center;'>Etelvia · Motor de Inteligencia Competitiva 360 · " + hoy + "</p></body></html>"
    return html

def generar_word_conjunto():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        t = doc.add_heading('Informe Estrategico Completo', 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(sector + " · " + tam + " · " + reg + " · " + hoy)
        doc.add_paragraph("SSG: " + str(ssg) + "/100 · Innovacion: " + str(macro_inn) + "/5")

        if inc_estrategico:
            doc.add_page_break()
            doc.add_heading('1. Informe Estrategico Competitivo', level=1)
            tb = doc.add_table(rows=1, cols=3); tb.style = 'Table Grid'
            h = tb.rows[0].cells; h[0].text='Indice'; h[1].text='Valor'; h[2].text='Nivel'
            for n,v in [('SSG',ssg),('ICE',ice),('ISF',isf),('IEO',ieo),('IDC',idc),('IIE',iie),('IPT',ipt)]:
                r=tb.add_row().cells; r[0].text=n; r[1].text=str(v)+'/100'; r[2].text=nivel(v)
            doc.add_heading('Variables Economicas', level=2)
            tb2=doc.add_table(rows=1,cols=3); tb2.style='Table Grid'
            h2=tb2.rows[0].cells; h2[0].text='Variable'; h2[1].text='Valor'; h2[2].text='Unidad'
            for n,v,u in [('Ventas',ventas,'k EUR'),('Empleados',empleados,''),('ROA',roa,'%'),('Endeudamiento',endeud,'')]:
                r=tb2.add_row().cells; r[0].text=n; r[1].text=str(round(v,2)); r[2].text=u

        if inc_innovacion:
            doc.add_page_break()
            doc.add_heading('2. Informe de Innovacion', level=1)
            doc.add_paragraph("Indice Global: " + str(macro_inn) + "/5 — " + nivel_inn)
            tb3=doc.add_table(rows=1,cols=3); tb3.style='Table Grid'
            h3=tb3.rows[0].cells; h3[0].text='Indicador'; h3[1].text='Punt.'; h3[2].text='Nivel'
            for n,v in [('I+D+i',b1),('Gest.Proyectos',b2),('Dllo.Productos',b3),('Estr.Innovacion',b4),('Desempeno',b5)]:
                r=tb3.add_row().cells; r[0].text=n; r[1].text=str(v)+'/5'; r[2].text=nivel5(v)

        if inc_benchmarking:
            doc.add_page_break()
            doc.add_heading('3. Benchmarking Estrategico', level=1)
            doc.add_paragraph("Cuadrante: " + get_cuadrante())
            tb4=doc.add_table(rows=1,cols=3); tb4.style='Table Grid'
            h4=tb4.rows[0].cells; h4[0].text='Indice'; h4[1].text='Valor'; h4[2].text='Nivel'
            for n,v in [('SSG',ssg),('ICE',ice),('ISF',isf),('IEO',ieo),('IDC',idc),('IIE',iie),('IPT',ipt)]:
                r=tb4.add_row().cells; r[0].text=n; r[1].text=str(v)+'/100'; r[2].text=nivel(v)

        if inc_plan:
            doc.add_page_break()
            doc.add_heading('4. Plan de Accion y Hoja de Ruta', level=1)
            doc.add_heading('Acciones Estrategicas', level=2)
            for trim,accion,impacto in get_acciones():
                doc.add_paragraph(trim + " · " + accion + " · Impacto: " + impacto)
            doc.add_heading('Hoja de Ruta 3 Anos', level=2)
            doc.add_paragraph("Ano 1: Consolidacion — Focalizar en brechas criticas.")
            doc.add_paragraph("Ano 2: Aceleracion — Escalar iniciativas de mayor impacto.")
            doc.add_paragraph("Ano 3: Liderazgo — Consolidar posicion de referencia.")
            doc.add_heading('Notas', level=2)
            for _ in range(6): doc.add_paragraph('_'*80)

        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)
        return buf.getvalue()
    except ImportError:
        return None

if not any([inc_estrategico, inc_innovacion, inc_benchmarking, inc_plan]):
    st.warning("Selecciona al menos un informe.")
else:
    html_out = generar_html_conjunto()
    word_out = generar_word_conjunto()
    c1,c2,c3 = st.columns(3)
    with c1:
        st.download_button("📄 Descargar Word completo", data=word_out if word_out else b"",
            file_name="informe_completo_" + hoy.replace('/','_') + ".docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary", use_container_width=True, disabled=word_out is None)
    with c2:
        st.download_button("🌐 Descargar HTML completo", data=html_out,
            file_name="informe_completo_" + hoy.replace('/','_') + ".html",
            mime="text/html", use_container_width=True)
    with c3:
        st.info("Para PDF: abre el HTML → **Ctrl+P** → **Guardar como PDF**")
    st.markdown("---")
    incluidos = []
    if inc_estrategico: incluidos.append("✅ Estratégico")
    if inc_innovacion:  incluidos.append("✅ Innovación")
    if inc_benchmarking: incluidos.append("✅ Benchmarking")
    if inc_plan: incluidos.append("✅ Plan de Acción")
    st.info("Informes incluidos: " + " · ".join(incluidos))