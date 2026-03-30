import streamlit as st
from asistentes import mostrar_felix
import pandas as pd
import plotly.graph_objects as go
import json
import os
import numpy as np
from streamlit_echarts import st_echarts

st.set_page_config(page_title="Índices Estratégicos", layout="wide")

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Rajdhani:wght@400;600;700&family=Inter:wght@300;400;500&display=swap');
.section-title {
    font-family: 'Rajdhani', sans-serif;
    font-size: 1.25rem; font-weight: 700; color: #ffffff;
    background: linear-gradient(90deg, rgba(124,58,237,0.25), rgba(124,58,237,0.0));
    border-left: 5px solid #7c3aed;
    padding: 10px 16px; border-radius: 0 8px 8px 0;
    margin: 20px 0 15px 0;
    text-shadow: 0 0 12px rgba(124,58,237,0.8);
}
.index-card {
    background: linear-gradient(135deg, #0d1b2a, #1a1035);
    border: 1px solid #2d1b69; border-radius: 14px;
    padding: 20px; text-align: center; position: relative; overflow: hidden;
}
.index-card::before {
    content:''; position:absolute; top:0; left:0; right:0; height:4px;
    background: linear-gradient(90deg,#7c3aed,#a855f7,#c084fc);
}
.index-label { font-size:0.68rem; color:#94a3b8; text-transform:uppercase; letter-spacing:1.5px; margin-bottom:4px; }
.index-name  { font-family:'Rajdhani',sans-serif; font-size:1.1rem; font-weight:700; color:#c084fc; margin-bottom:8px; }
.index-value { font-family:'Rajdhani',sans-serif; font-size:3rem; font-weight:700; color:#a855f7; line-height:1; }
.index-pct   { color:#7c3aed; font-size:0.8rem; margin-top:4px; }
.index-desc  { color:#64748b; font-size:0.7rem; margin-top:6px; line-height:1.4; }
.ssg-card {
    background: linear-gradient(135deg, #0a0e1a, #1a0a2e);
    border: 2px solid #7c3aed; border-radius: 16px;
    padding: 28px; text-align: center; position: relative; overflow: hidden;
}
.ssg-card::before {
    content:''; position:absolute; top:0; left:0; right:0; height:5px;
    background: linear-gradient(90deg,#00d4ff,#7c3aed,#f59e0b,#10b981);
}
.badge {
    display:inline-block; padding:3px 10px; border-radius:20px;
    font-size:0.72rem; font-weight:600; margin-top:6px;
}
.badge-alto    { background:rgba(16,185,129,0.2); color:#10b981; border:1px solid #10b981; }
.badge-medio   { background:rgba(245,158,11,0.2); color:#f59e0b; border:1px solid #f59e0b; }
.badge-bajo    { background:rgba(239,68,68,0.2);  color:#ef4444; border:1px solid #ef4444; }
</style>
""", unsafe_allow_html=True)

# ── Cargar perfil ─────────────────────────────────────────────────────────────
PERFIL_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'perfil_empresa.json')

def cargar_perfil():
    if os.path.exists(PERFIL_FILE):
        with open(PERFIL_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

perfil = cargar_perfil()
for k, v in perfil.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ── Cargar datos de empresa desde Supabase si no están en session_state ───────
if not st.session_state.get('save_reg_user'):
    try:
        from supabase import create_client
        _url = st.secrets.get("SUPABASE_URL", os.environ.get("SUPABASE_URL",""))
        _key = st.secrets.get("SUPABASE_KEY", os.environ.get("SUPABASE_KEY",""))
        _ec  = st.session_state.get('empresa_codigo') or cargar_perfil().get('sesion_codigo','')
        if _url and _key and _ec:
            _sb = create_client(_url, _key)
            _e  = _sb.table('empresas').select('*').eq('codigo',_ec).execute().data
            if _e:
                _e = _e[0]
                _SM = {"Alimentación y bebidas":1,"Textil y confección":2,"Cuero y calzado":3,"Química y plásticos":4,"Minerales no metálicos":5,"Metalmecanico":6,"Maquinaria equipo":7,"Otras manufacturas":8,"Electrónica, telecomunicaciones":9,"Informática, software. Robótica, IA":10,"Actividades I+D: biotech, farmacia":11,"Transporte y logística":12,"Consultoría y servicios profesionales":13,"Turismo y hosteleria":14,"Retail y comercio":15,"Otros servicios":16}
                _MM = {1:1,2:1,3:1,4:1,5:1,6:1,7:1,8:1,9:2,10:2,11:2,12:3,13:3,14:3,15:3,16:3}
                _sc = _SM.get(_e.get('sector',''),0)
                st.session_state.update({
                    'save_sector_nombre':_e.get('sector',''),'save_tam_nombre':_e.get('tamano',''),
                    'save_reg_nombre':_e.get('region',''),'save_export_nombre':_e.get('exportacion',''),
                    'save_anti_nombre':_e.get('antiguedad',''),'save_ventas':float(_e.get('ventas') or 0),
                    'save_empleados':int(_e.get('empleados') or 0),'save_roa':float(_e.get('roa') or 0),
                    'save_var_vtas':float(_e.get('var_ventas') or 0),'save_var_empl':float(_e.get('var_empleados') or 0),
                    'save_productiv':float(_e.get('productividad') or 0),'save_coste_emp':float(_e.get('coste_empleado') or 0),
                    'save_endeud':float(_e.get('endeudamiento') or 0),
                    'save_reg_user':{"Andalucia":1,"Aragon":2,"Asturias":3,"Baleares":4,"Canarias":5,"Cantabria":6,"Castilla la Mancha":7,"Castilla y León":8,"Cataluña":9,"Com Valenciana":10,"Extremadura":11,"Galicia":12,"Madrid":13,"Murcia":14,"Navarra":15,"Pais Vasco":16}.get(_e.get('region',''),0),
                    'save_tam_user':{"Pequeña":1,"Mediana":2,"Grande":3}.get(_e.get('tamano',''),0),
                    'save_sector_cod':_sc,'save_macro_cod':_MM.get(_sc,3),
                    'save_export_cod':{"Menos 10 %":1,"10 - 30 %":2,"30 - 60 %":3,"> 60 %":4}.get(_e.get('exportacion',''),0),
                    'save_anti_cod':{"Menos 10 años":1,"10-30 años":2,"> 30 años":3}.get(_e.get('antiguedad',''),0),
                })
    except Exception: pass

if not st.session_state.get('save_reg_user'):
    st.warning("⚠️ Primero completa tu perfil de empresa en **Mi Empresa → Datos de la empresa**.")
    st.stop()

# ── Cargar Excel ──────────────────────────────────────────────────────────────
ruta_excel = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'datos.xlsx')
if not os.path.exists(ruta_excel): ruta_excel = "datos.xlsx"
try:
    df = pd.read_excel(ruta_excel)
    df.columns = df.columns.str.strip()
except Exception as e:
    st.error(f"Error al cargar datos: {e}")
    st.stop()

# ── Mapas ─────────────────────────────────────────────────────────────────────
REGIONES     = {1:"Andalucia",2:"Aragon",3:"Asturias",4:"Baleares",5:"Canarias",6:"Cantabria",7:"Castilla la Mancha",8:"Castilla y León",9:"Cataluña",10:"Com Valenciana",11:"Extremadura",12:"Galicia",13:"Madrid",14:"Murcia",15:"Navarra",16:"Pais Vasco"}
TAMANIOS     = {1:"Pequeña",2:"Mediana",3:"Grande"}
MACROSECTORES= {1:"Industria",2:"Tecnologia avanzada",3:"Servicios"}
EXPORTACIONES= {1:"Menos 10 %",2:"10 - 30 %",3:"30 - 60 %",4:"> 60 %"}
ANTIGUEDADES = {1:"Menos 10 años",2:"10-30 años",3:"> 30 años"}
SECTORES     = {1:"Alimentación y bebidas",2:"Textil y confección",3:"Cuero y calzado",4:"Química y plásticos",5:"Minerales no metálicos",6:"Metalmecánico",7:"Maquinaria equipo",8:"Otras manufacturas",9:"Electrónica, telecomunicaciones",10:"Informática, software. Robótica, IA",11:"Actividades I+D: biotech, farmacia",12:"Transporte y logística",13:"Consultoría y servicios profesionales",14:"Turismo y hosteleria",15:"Retail y comercio",16:"Otros servicios"}
SECTOR_MACRO = {1:1,2:1,3:1,4:1,5:1,6:1,7:1,8:1,9:2,10:2,11:2,12:3,13:3,14:3,15:3,16:3}

# ═════════════════════════════════════════════════════════════════════════════
# FUNCIÓN: CALCULAR ÍNDICES SOBRE UN DATAFRAME
# ═════════════════════════════════════════════════════════════════════════════
def pct_rank(series, valor, invertido=False):
    if invertido:
        return float(np.sum(series >= valor) / len(series) * 100)
    return float(np.sum(series <= valor) / len(series) * 100)

def norm_col(series, valor, invertido=False):
    return pct_rank(series, valor, invertido)

def calcular_indices(df_ref, mis_datos):
    def p(col, inv=False):
        if col not in df_ref.columns: return 50.0
        return norm_col(df_ref[col], mis_datos.get(col, df_ref[col].mean()), inv)

    ICE = (p('Prod_Venta_Emp')*0.30 + p('Var_Ventas_5a')*0.25 + p('Var_Emp_5a')*0.15 +
           p('Ratio_Endeudamiento',True)*0.15 + p('Exportacion')*0.15)
    ISF = (p('Ratio_Endeudamiento',True)*0.40 + p('ROA')*0.35 + p('Coste_Med_Emp',True)*0.25)
    IEO = (p('Prod_Venta_Emp')*0.50 + p('Coste_Med_Emp',True)*0.30 + p('Var_Ventas_5a')*0.20)
    IDC = (p('Var_Ventas_5a')*0.40 + p('Var_Emp_5a')*0.35 + p('ROA')*0.25)
    IIE = (p('Exportacion')*0.60 + p('Var_Ventas_5a')*0.20 + p('Prod_Venta_Emp')*0.20)
    IPT = (p('Prod_Venta_Emp')*0.40 + p('Coste_Med_Emp')*0.30 + p('Var_Emp_5a')*0.30)
    SSG = (ICE*0.25 + ISF*0.20 + IEO*0.20 + IDC*0.15 + IIE*0.10 + IPT*0.10)

    return {'ICE':round(ICE,1),'ISF':round(ISF,1),'IEO':round(IEO,1),
            'IDC':round(IDC,1),'IIE':round(IIE,1),'IPT':round(IPT,1),'SSG':round(SSG,1)}

def calcular_indices_df(df_ref, df_calc):
    resultados = []
    for _, row in df_calc.iterrows():
        resultados.append(calcular_indices(df_ref, row.to_dict()))
    return pd.DataFrame(resultados)

# ── Datos de mi empresa ───────────────────────────────────────────────────────
mis_datos = {
    'Ventas':              st.session_state.get('save_ventas', 0),
    'Empleados':           st.session_state.get('save_empleados', 0),
    'ROA':                 st.session_state.get('save_roa', 0),
    'Var_Ventas_5a':       st.session_state.get('save_var_vtas', 0),
    'Var_Emp_5a':          st.session_state.get('save_var_empl', 0),
    'Prod_Venta_Emp':      st.session_state.get('save_productiv', 0),
    'Coste_Med_Emp':       st.session_state.get('save_coste_emp', 0),
    'Ratio_Endeudamiento': st.session_state.get('save_endeud', 0),
    'Exportacion':         st.session_state.get('save_export_cod', 1),
    'Region':              st.session_state.get('save_reg_user', 1),
    'Tamaño':              st.session_state.get('save_tam_user', 1),
    'Macrosector':         st.session_state.get('save_macro_cod', 1),
    'Sector':              st.session_state.get('save_sector_cod', 1),
    'Antigüedad':          st.session_state.get('save_anti_cod', 1),
}

# ═════════════════════════════════════════════════════════════════════════════
# SIDEBAR: GRUPO DE REFERENCIA
# ═════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("### 🔍 Grupo de Referencia")
    st.caption("Elige con qué empresas calcular tu posición. Vacío = todas.")
    st.divider()
    REG_INV = {v:k for k,v in REGIONES.items()}
    TAM_INV = {v:k for k,v in TAMANIOS.items()}
    MAC_INV = {v:k for k,v in MACROSECTORES.items()}
    EXP_INV = {v:k for k,v in EXPORTACIONES.items()}
    ANT_INV = {v:k for k,v in ANTIGUEDADES.items()}
    SEC_INV = {v:k for k,v in SECTORES.items()}

    mi_reg = REGIONES.get(st.session_state.get('save_reg_user'))
    mi_tam = TAMANIOS.get(st.session_state.get('save_tam_user'))

    reg_sel = st.multiselect("📍 Región",     list(REGIONES.values()),     default=[mi_reg] if mi_reg else [])
    tam_sel = st.multiselect("🏢 Tamaño",     list(TAMANIOS.values()),     default=[mi_tam] if mi_tam else [])
    mac_sel = st.multiselect("🏭 Macrosector",list(MACROSECTORES.values()),default=[])
    if mac_sel:
        sects_disp = [v for k,v in SECTORES.items() if SECTOR_MACRO.get(k) in [MAC_INV.get(m) for m in mac_sel]]
    else:
        sects_disp = list(SECTORES.values())
    sec_sel = st.multiselect("🔧 Sector",     sects_disp,                 default=[])
    exp_sel = st.multiselect("🌍 Exportación",list(EXPORTACIONES.values()),default=[])
    ant_sel = st.multiselect("📅 Antigüedad", list(ANTIGUEDADES.values()), default=[])

# ── Aplicar filtros ───────────────────────────────────────────────────────────
dff = df.copy()
if reg_sel: dff = dff[dff['Region'].isin([REG_INV[r] for r in reg_sel if r in REG_INV])]
if tam_sel: dff = dff[dff['Tamaño'].isin([TAM_INV[t] for t in tam_sel if t in TAM_INV])]
if mac_sel: dff = dff[dff['Macrosector'].isin([MAC_INV[m] for m in mac_sel if m in MAC_INV])]
if sec_sel: dff = dff[dff['Sector'].isin([SEC_INV[s] for s in sec_sel if s in SEC_INV])]
if exp_sel: dff = dff[dff['Exportacion'].isin([EXP_INV[e] for e in exp_sel if e in EXP_INV])]
if ant_sel: dff = dff[dff['Antigüedad'].isin([ANT_INV[a] for a in ant_sel if a in ANT_INV])]

n = len(dff)

# ── Calcular índices ──────────────────────────────────────────────────────────
idx_empresa  = calcular_indices(dff, mis_datos)
idx_total    = calcular_indices(df,  mis_datos)
grp_indices  = calcular_indices_df(dff, dff)
grp_medias   = {k: round(grp_indices[k].mean(), 1) for k in ['ICE','ISF','IEO','IDC','IIE','IPT','SSG']}
tot_indices  = calcular_indices_df(df, df)
tot_medias   = {k: round(tot_indices[k].mean(), 1) for k in ['ICE','ISF','IEO','IDC','IIE','IPT','SSG']}

INDICES_INFO = {
    'ICE': ('Competitividad Empresarial',  'Productividad + crecimiento + apertura exterior',    '🏆'),
    'ISF': ('Solidez Financiera',          'Endeudamiento + rentabilidad + eficiencia de coste', '🛡️'),
    'IEO': ('Eficiencia Operativa',        'Productividad + control costes + crecimiento',       '⚙️'),
    'IDC': ('Dinamismo y Crecimiento',     'Crecimiento ventas + empleo + rentabilidad',         '🚀'),
    'IIE': ('Intensidad Exportadora',      'Apertura internacional + crecimiento + productividad','🌍'),
    'IPT': ('Productividad y Talento',     'Productividad + calidad del empleo + expansión',     '💡'),
}

def badge(v):
    if v >= 66: return '<span class="badge badge-alto">Alto</span>'
    if v >= 33: return '<span class="badge badge-medio">Medio</span>'
    return '<span class="badge badge-bajo">Bajo</span>'

def color_val(v):
    if v >= 66: return '#10b981'
    if v >= 33: return '#f59e0b'
    return '#ef4444'

# ═════════════════════════════════════════════════════════════════════════════
# CABECERA
# ═════════════════════════════════════════════════════════════════════════════
nom_reg = st.session_state.get('save_reg_nombre','—')
nom_tam = st.session_state.get('save_tam_nombre','—')
ssg_v   = idx_empresa['SSG']
ssg_col = color_val(ssg_v)

st.markdown(f"""
<div style='background:linear-gradient(135deg,#0a0e1a,#1a1035);border-radius:14px;padding:22px 28px;margin-bottom:18px;border:1px solid #2d1b69;'>
    <h1 style='font-family:Rajdhani,sans-serif;color:#a855f7;margin:0;font-size:1.9rem;'>📐 Índices Estratégicos y de Competitividad</h1>
    <p style='color:#94a3b8;margin:6px 0 0 0;font-size:0.88rem;'>
        {st.session_state.get('save_sector_nombre','Tu empresa')} · 
        Región: <strong style='color:#e2e8f0;'>{nom_reg}</strong> · 
        Tamaño: <strong style='color:#e2e8f0;'>{nom_tam}</strong> · 
        Comparando con <strong style='color:#a855f7;'>{n} empresas</strong>
    </p>
</div>
""", unsafe_allow_html=True)

if n < 10:
    st.error(f"⚠️ Solo {n} empresas en el grupo. Necesitas al menos 10. Reduce los filtros.")
    st.stop()

# ═════════════════════════════════════════════════════════════════════════════
# SCORE ESTRATÉGICO GLOBAL (SSG)
# ═════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="section-title">⭐ Score Estratégico Global (SSG)</div>', unsafe_allow_html=True)

c_ssg1, c_ssg2, c_ssg3 = st.columns([1,1,2])
with c_ssg1:
    st.markdown(f"""
    <div class="ssg-card">
        <div style="font-size:0.75rem;color:#64748b;text-transform:uppercase;letter-spacing:2px;">Score Estratégico Global</div>
        <div style="font-family:Rajdhani,sans-serif;font-size:5rem;font-weight:700;color:{ssg_col};line-height:1;margin:8px 0;">
            {ssg_v}
        </div>
        <div style="color:#94a3b8;font-size:0.85rem;">sobre 100</div>
        {badge(ssg_v)}
    </div>
    """, unsafe_allow_html=True)

with c_ssg2:
    fig_ssg = go.Figure(go.Indicator(
        mode="gauge+number",
        value=ssg_v,
        number={'font': {'size': 36, 'color': ssg_col}},
        title={'text': "SSG", 'font': {'size': 13, 'color': '#94a3b8'}},
        gauge={
            'axis': {'range': [0, 100], 'tickcolor': '#334155', 'tickfont': {'color': '#64748b', 'size': 9}},
            'bar': {'color': ssg_col, 'thickness': 0.28},
            'bgcolor': '#0a0e1a', 'borderwidth': 0,
            'steps': [
                {'range': [0,  33], 'color': '#1a0505'},
                {'range': [33, 66], 'color': '#1a1200'},
                {'range': [66,100], 'color': '#051a0a'},
            ],
            'threshold': {'line': {'color': '#7c3aed', 'width': 3}, 'thickness': 0.75, 'value': grp_medias['SSG']}
        }
    ))
    fig_ssg.update_layout(height=220, margin=dict(l=20,r=20,t=40,b=10),
        paper_bgcolor='rgba(0,0,0,0)', font_color='#94a3b8')
    st.plotly_chart(fig_ssg, use_container_width=True)
    st.caption(f"🟣 Línea = media grupo ({grp_medias['SSG']})")

with c_ssg3:
    pesos_ssg = {'ICE':0.25,'ISF':0.20,'IEO':0.20,'IDC':0.15,'IIE':0.10,'IPT':0.10}
    contrib_labels = list(pesos_ssg.keys())
    contrib_vals   = [round(idx_empresa[k]*v, 1) for k,v in pesos_ssg.items()]
    contrib_cols   = [color_val(idx_empresa[k]) for k in pesos_ssg.keys()]

    ssg_comp = {
        "backgroundColor": "#0a0e1a",
        "tooltip": {"trigger":"axis","axisPointer":{"type":"shadow"},
                    "backgroundColor":"#1e293b","textStyle":{"color":"#e2e8f0"}},
        "title": {"text": "Contribución de cada índice al SSG",
                  "textStyle": {"color": "#94a3b8", "fontSize": 11}, "top": 5},
        "grid": {"left":"5%","right":"12%","bottom":"5%","containLabel":True},
        "xAxis": {"type":"value","max":25,"axisLabel":{"color":"#64748b"},
                  "splitLine":{"lineStyle":{"color":"#1e3a5f"}}},
        "yAxis": {"type":"category","data":contrib_labels,"inverse":True,
                  "axisLabel":{"color":"#94a3b8","fontSize":12}},
        "series":[{
            "type":"bar",
            "data":[{"value":v,"itemStyle":{"color":c,"borderRadius":[0,4,4,0]}}
                    for v,c in zip(contrib_vals, contrib_cols)],
            "label":{"show":True,"position":"right","color":"#e2e8f0","fontSize":11,
                     "formatter":"{c}"},
        }]
    }
    st_echarts(options=ssg_comp, height="220px")

st.divider()

# ═════════════════════════════════════════════════════════════════════════════
# CARDS DE LOS 6 ÍNDICES
# ═════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="section-title">📐 Los 6 Índices Estratégicos</div>', unsafe_allow_html=True)

card_cols = st.columns(6)
for i, (cod, (nombre, desc, icono)) in enumerate(INDICES_INFO.items()):
    v    = idx_empresa[cod]
    gv   = grp_medias[cod]
    tv   = tot_medias[cod]
    diff = round(v - gv, 1)
    diff_icon = "▲" if diff >= 0 else "▼"
    diff_col  = "#10b981" if diff >= 0 else "#ef4444"
    card_cols[i].markdown(f"""
    <div class="index-card">
        <div class="index-label">{icono} {cod}</div>
        <div class="index-name">{nombre}</div>
        <div class="index-value" style="color:{color_val(v)};">{v}</div>
        <div style="color:{diff_col};font-size:0.78rem;">{diff_icon} {abs(diff)} vs grupo</div>
        <div class="index-pct">Media grupo: {gv} · Total: {tv}</div>
        {badge(v)}
        <div class="index-desc">{desc}</div>
    </div>""", unsafe_allow_html=True)

st.divider()

# ═════════════════════════════════════════════════════════════════════════════
# RADAR INTERACTIVO
# ═════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="section-title">🕸️ Radar de Índices — Mi empresa vs. grupo vs. total</div>', unsafe_allow_html=True)

c_rad, c_bar = st.columns(2)
indices_labels = list(INDICES_INFO.keys())
emp_r = [idx_empresa[k] for k in indices_labels]
grp_r = [grp_medias[k]  for k in indices_labels]
tot_r = [tot_medias[k]  for k in indices_labels]

with c_rad:
    radar_opt = {
        "backgroundColor": "#0a0e1a",
        "tooltip": {"trigger":"item","backgroundColor":"#1e293b","borderColor":"#2d1b69",
                    "textStyle":{"color":"#e2e8f0"}},
        "legend": {"data":["Mi Empresa", f"Media Grupo ({n})", "Media Total"],
                   "textStyle":{"color":"#94a3b8"},"bottom":0},
        "radar": {
            "indicator": [{"name":k,"max":100} for k in indices_labels],
            "shape":"polygon","splitNumber":4,
            "axisName":{"color":"#94a3b8","fontSize":12,"fontFamily":"Rajdhani"},
            "splitLine":{"lineStyle":{"color":["#2d1b69","#2d1b69","#2d1b69","#2d1b69"]}},
            "splitArea":{"areaStyle":{"color":["rgba(124,58,237,0.03)","rgba(124,58,237,0.07)"]}},
            "axisLine":{"lineStyle":{"color":"#2d1b69"}},
        },
        "series":[{"type":"radar","data":[
            {"value":emp_r,"name":"Mi Empresa",
             "itemStyle":{"color":"#a855f7"},"lineStyle":{"color":"#a855f7","width":2},
             "areaStyle":{"color":"rgba(168,85,247,0.2)"},"symbol":"circle","symbolSize":6},
            {"value":grp_r,"name":f"Media Grupo ({n})",
             "itemStyle":{"color":"#10b981"},"lineStyle":{"color":"#10b981","width":2,"type":"dashed"},
             "areaStyle":{"color":"rgba(16,185,129,0.07)"},"symbol":"circle","symbolSize":5},
            {"value":tot_r,"name":"Media Total",
             "itemStyle":{"color":"#f59e0b"},"lineStyle":{"color":"#f59e0b","width":1.5,"type":"dotted"},
             "areaStyle":{"color":"rgba(245,158,11,0.04)"},"symbol":"circle","symbolSize":4},
        ]}],
    }
    st_echarts(options=radar_opt, height="420px")

with c_bar:
    bar_opt = {
        "backgroundColor": "#0a0e1a",
        "tooltip": {"trigger":"axis","axisPointer":{"type":"shadow"},
                    "backgroundColor":"#1e293b","textStyle":{"color":"#e2e8f0"}},
        "legend": {"data":["Mi Empresa", f"Grupo ({n})", "Total"],
                   "textStyle":{"color":"#94a3b8"},"top":5},
        "grid": {"left":"3%","right":"4%","bottom":"10%","containLabel":True},
        "xAxis": {"type":"category","data":indices_labels,
                  "axisLabel":{"color":"#94a3b8","fontSize":12},
                  "axisLine":{"lineStyle":{"color":"#2d1b69"}}},
        "yAxis": {"type":"value","min":0,"max":100,
                  "axisLabel":{"color":"#64748b"},
                  "splitLine":{"lineStyle":{"color":"#1e3a5f"}}},
        "series":[
            {"name":"Mi Empresa","type":"bar","data":emp_r,
             "itemStyle":{"color":"#a855f7","borderRadius":[4,4,0,0]},
             "label":{"show":True,"position":"top","color":"#a855f7","fontSize":10,"formatter":"{c}"}},
            {"name":f"Grupo ({n})","type":"bar","data":grp_r,
             "itemStyle":{"color":"#10b981","borderRadius":[4,4,0,0]},
             "label":{"show":True,"position":"top","color":"#10b981","fontSize":10,"formatter":"{c}"}},
            {"name":"Total","type":"bar","data":tot_r,
             "itemStyle":{"color":"#f59e0b","borderRadius":[4,4,0,0]},
             "label":{"show":True,"position":"top","color":"#f59e0b","fontSize":10,"formatter":"{c}"}},
        ]
    }
    st_echarts(options=bar_opt, height="420px")

st.divider()

# ═════════════════════════════════════════════════════════════════════════════
# VELOCÍMETROS
# ═════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="section-title">⚡ Velocímetros — Posición en cada índice</div>', unsafe_allow_html=True)

g_cols = st.columns(6)
for i, (cod, (nombre, desc, icono)) in enumerate(INDICES_INFO.items()):
    v  = idx_empresa[cod]
    gv = grp_medias[cod]
    with g_cols[i]:
        fig_g = go.Figure(go.Indicator(
            mode="gauge+number+delta",
            value=v,
            delta={'reference': gv, 'valueformat': '.1f',
                   'increasing': {'color':'#10b981'}, 'decreasing': {'color':'#ef4444'}},
            number={'valueformat':'.1f','font':{'size':22,'color':color_val(v)}},
            title={'text': f"{icono} {cod}", 'font':{'size':11,'color':'#94a3b8'}},
            gauge={
                'axis':{'range':[0,100],'tickcolor':'#334155','tickfont':{'color':'#64748b','size':7}},
                'bar':{'color':color_val(v),'thickness':0.25},
                'bgcolor':'#0a0e1a','borderwidth':0,
                'steps':[
                    {'range':[0, 33],'color':'#1a0505'},
                    {'range':[33,66],'color':'#1a1200'},
                    {'range':[66,100],'color':'#051a0a'},
                ],
                'threshold':{'line':{'color':'#7c3aed','width':3},'thickness':0.75,'value':gv}
            }
        ))
        fig_g.update_layout(height=190, margin=dict(l=10,r=10,t=35,b=5),
            paper_bgcolor='rgba(0,0,0,0)', font_color='#94a3b8')
        st.plotly_chart(fig_g, use_container_width=True)
        st.caption(f"🟣 Media grupo: {gv}")

st.divider()

# ═════════════════════════════════════════════════════════════════════════════
# RANKING PERCENTIL
# ═════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="section-title">🏆 Ranking Percentil de los 6 Índices</div>', unsafe_allow_html=True)

pct_vals  = [idx_empresa[k] for k in indices_labels]
pct_cols  = [color_val(v) for v in pct_vals]
pct_names = [f"{icono} {cod} — {nombre}" for cod, (nombre, _, icono) in INDICES_INFO.items()]

rank_opt = {
    "backgroundColor": "#0a0e1a",
    "tooltip": {"trigger":"axis","axisPointer":{"type":"shadow"},
                "backgroundColor":"#1e293b","textStyle":{"color":"#e2e8f0"},
                "formatter":"{b}: {c} / 100"},
    "grid": {"left":"2%","right":"10%","bottom":"5%","containLabel":True},
    "xAxis": {"type":"value","min":0,"max":100,
              "axisLabel":{"color":"#64748b","formatter":"{value}"},
              "splitLine":{"lineStyle":{"color":"#1e3a5f"}}},
    "yAxis": {"type":"category","data":pct_names,"inverse":True,
              "axisLabel":{"color":"#94a3b8","fontSize":11}},
    "series":[{
        "type":"bar",
        "data":[{"value":v,"itemStyle":{"color":c,"borderRadius":[0,4,4,0]}}
                for v,c in zip(pct_vals, pct_cols)],
        "label":{"show":True,"position":"right","color":"#e2e8f0","fontSize":12,
                 "fontFamily":"Rajdhani","formatter":"{c} / 100"},
        "markLine":{"silent":True,"lineStyle":{"color":"#7c3aed","type":"dashed","width":2},
            "data":[{"xAxis":50,"label":{"formatter":"Mediana","color":"#7c3aed","fontSize":10}}]},
    }]
}
st_echarts(options=rank_opt, height="380px")

st.divider()

# ═════════════════════════════════════════════════════════════════════════════
# RANKINGS POR PERFIL
# ═════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="section-title">📊 Rankings Comparativos por Perfil</div>', unsafe_allow_html=True)

tab1, tab2, tab3, tab4 = st.tabs(["🏭 Por Macrosector", "📍 Por Región", "🏢 Por Tamaño", "🌍 Por Exportación"])

def ranking_tab(df_ref, col_agrup, mapa):
    ind_sel_rank = st.selectbox("Índice a comparar", list(INDICES_INFO.keys()),
                                key=f"rank_{col_agrup}", index=0)
    rows = []
    for cod, nombre in mapa.items():
        subset = df_ref[df_ref[col_agrup] == cod]
        if len(subset) >= 5:
            idx_sub = calcular_indices_df(df_ref, subset)
            rows.append({'Grupo': nombre, 'N': len(subset),
                         ind_sel_rank: round(idx_sub[ind_sel_rank].mean(), 1)})
    if not rows:
        st.info("No hay suficientes datos para este ranking.")
        return
    df_rank = pd.DataFrame(rows).sort_values(ind_sel_rank, ascending=False).reset_index(drop=True)
    bar_data   = df_rank[ind_sel_rank].tolist()
    bar_labels = df_rank['Grupo'].tolist()
    bar_colors = [color_val(v) for v in bar_data]
    opt = {
        "backgroundColor": "#0a0e1a",
        "tooltip": {"trigger":"axis","backgroundColor":"#1e293b","textStyle":{"color":"#e2e8f0"}},
        "grid": {"left":"3%","right":"10%","bottom":"5%","containLabel":True},
        "xAxis": {"type":"value","min":0,"max":100,"axisLabel":{"color":"#64748b"},
                  "splitLine":{"lineStyle":{"color":"#1e3a5f"}}},
        "yAxis": {"type":"category","data":bar_labels,"inverse":True,
                  "axisLabel":{"color":"#94a3b8","fontSize":11}},
        "series":[{"type":"bar",
            "data":[{"value":v,"itemStyle":{"color":c,"borderRadius":[0,4,4,0]}}
                    for v,c in zip(bar_data, bar_colors)],
            "label":{"show":True,"position":"right","color":"#e2e8f0","fontSize":11,"formatter":"{c}"}}]
    }
    st_echarts(options=opt, height=f"{max(250, len(bar_labels)*45)}px")
    st.dataframe(df_rank.rename(columns={ind_sel_rank: f"Media {ind_sel_rank}"}),
                 use_container_width=True, hide_index=True)

with tab1: ranking_tab(dff, 'Macrosector', MACROSECTORES)
with tab2: ranking_tab(dff, 'Region',      REGIONES)
with tab3: ranking_tab(dff, 'Tamaño',      TAMANIOS)
with tab4: ranking_tab(dff, 'Exportacion', EXPORTACIONES)

st.divider()

# ═════════════════════════════════════════════════════════════════════════════
# TABLA RESUMEN
# ═════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="section-title">📋 Tabla Resumen de Índices</div>', unsafe_allow_html=True)

tabla_data = []
for cod, (nombre, desc, icono) in INDICES_INFO.items():
    v  = idx_empresa[cod]
    gv = grp_medias[cod]
    tv = tot_medias[cod]
    nivel = "🟢 Alto" if v >= 66 else ("🟡 Medio" if v >= 33 else "🔴 Bajo")
    tabla_data.append({'Índice': f"{icono} {cod}",'Nombre': nombre,'Mi Empresa': v,
        f'Media Grupo ({n} emp)': gv,'Media Total (1.000)': tv,
        'Diferencia vs grupo': f"{v-gv:+.1f}",'Nivel': nivel})
tabla_data.append({'Índice':'⭐ SSG','Nombre':'Score Estratégico Global',
    'Mi Empresa':idx_empresa['SSG'],f'Media Grupo ({n} emp)':grp_medias['SSG'],
    'Media Total (1.000)':tot_medias['SSG'],
    'Diferencia vs grupo':f"{idx_empresa['SSG']-grp_medias['SSG']:+.1f}",
    'Nivel':"🟢 Alto" if idx_empresa['SSG']>=66 else ("🟡 Medio" if idx_empresa['SSG']>=33 else "🔴 Bajo")})
st.dataframe(pd.DataFrame(tabla_data), use_container_width=True, hide_index=True)

st.divider()
st.divider()

# ═════════════════════════════════════════════════════════════════════════════
# DESCARGA — HTML y WORD
# ═════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="section-title">📥 Descargar Índices Estratégicos</div>', unsafe_allow_html=True)

from datetime import date as _date
hoy = _date.today().strftime("%d/%m/%Y")
nom_sector = st.session_state.get('save_sector_nombre','—')

def generar_html_indices():
    idx_tabla = "".join([
        f"<tr><td><strong>{icono} {cod}</strong></td><td>{nombre}</td>"
        f"<td style='text-align:center;font-weight:700;color:{'green' if idx_empresa[cod]>=66 else 'orange' if idx_empresa[cod]>=33 else 'red'};'>{idx_empresa[cod]}/100</td>"
        f"<td style='text-align:center;color:#666;'>{grp_medias[cod]}/100</td>"
        f"<td style='text-align:center;color:#888;'>{tot_medias[cod]}/100</td>"
        f"<td style='text-align:center;font-weight:600;color:{'green' if idx_empresa[cod]>=grp_medias[cod] else 'red'};'>{idx_empresa[cod]-grp_medias[cod]:+.1f}</td></tr>"
        for cod,(nombre,desc,icono) in INDICES_INFO.items()
    ])
    ssg_color = 'green' if idx_empresa['SSG']>=66 else 'orange' if idx_empresa['SSG']>=33 else 'red'
    return f"""<!DOCTYPE html><html><head><meta charset='utf-8'>
<style>body{{font-family:Georgia,serif;margin:50px auto;max-width:900px;color:#1a1a1a;line-height:1.75;}}
h1{{color:#7c3aed;border-bottom:3px solid #7c3aed;padding-bottom:10px;font-size:1.8rem;}}
.perfil{{background:#f5f3ff;border-radius:8px;padding:16px;margin:16px 0;font-size:.9rem;}}
.ssg{{font-size:2.8rem;font-weight:700;color:{ssg_color};}}
table{{border-collapse:collapse;width:100%;margin:16px 0;font-size:.9rem;}}
th{{background:#7c3aed;color:white;padding:10px;text-align:left;}}
td{{padding:9px 11px;border-bottom:1px solid #e5e7eb;}}
.notas{{border:1px dashed #9ca3af;border-radius:8px;padding:40px 16px;margin:20px 0;
    color:#9ca3af;font-style:italic;text-align:center;}}
@media print{{body{{margin:20px;}}}}</style></head><body>
<h1>📐 Índices Estratégicos y de Competitividad</h1>
<div class="perfil">
<strong>{nom_sector}</strong> | {st.session_state.get('save_tam_nombre','—')} | {st.session_state.get('save_reg_nombre','—')}<br>
Grupo de referencia: <strong>{n} empresas</strong> | Fecha: {hoy}
</div>
<p>Score Estratégico Global (SSG): <span class="ssg">{idx_empresa['SSG']}</span>/100
&nbsp;·&nbsp; {'🟢 Alto' if idx_empresa['SSG']>=66 else '🟡 Medio' if idx_empresa['SSG']>=33 else '🔴 Bajo'}</p>
<h2 style='color:#7c3aed;'>Los 6 Índices Estratégicos</h2>
<table>
<tr><th>Índice</th><th>Nombre</th><th>Mi Empresa</th><th>Media Grupo</th><th>Media Total</th><th>Dif. vs grupo</th></tr>
{idx_tabla}
</table>
<div class="notas">✏️ Espacio para tus anotaciones y comentarios</div>
<hr style='margin-top:40px;'/>
<p style='color:#9ca3af;font-size:.78rem;text-align:center;'>Plataforma de Diagnóstico Estratégico · {hoy}</p>
</body></html>"""

def generar_word_indices():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        doc = Document()
        t = doc.add_heading('Índices Estratégicos y de Competitividad', 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Empresa: {nom_sector} | {st.session_state.get('save_tam_nombre','—')} | {st.session_state.get('save_reg_nombre','—')}")
        doc.add_paragraph(f"Grupo de referencia: {n} empresas | Fecha: {hoy}")
        doc.add_paragraph(f"Score Estratégico Global (SSG): {idx_empresa['SSG']}/100 — {'Alto' if idx_empresa['SSG']>=66 else 'Medio' if idx_empresa['SSG']>=33 else 'Bajo'}")
        doc.add_heading('Los 6 Índices Estratégicos', level=1)
        tabla = doc.add_table(rows=1, cols=5); tabla.style = 'Table Grid'
        hdr = tabla.rows[0].cells
        hdr[0].text='Índice'; hdr[1].text='Nombre'; hdr[2].text='Mi Empresa'; hdr[3].text='Media Grupo'; hdr[4].text='Dif. vs grupo'
        for cod,(nombre,desc,icono) in INDICES_INFO.items():
            row = tabla.add_row().cells
            row[0].text=f"{icono} {cod}"; row[1].text=nombre
            row[2].text=f"{idx_empresa[cod]}/100"; row[3].text=f"{grp_medias[cod]}/100"
            row[4].text=f"{idx_empresa[cod]-grp_medias[cod]:+.1f}"
        doc.add_heading('Notas y comentarios del equipo directivo', level=1)
        doc.add_paragraph("[ Escribe aquí tus conclusiones y próximos pasos ]")
        for _ in range(8):
            p=doc.add_paragraph(); p.add_run('_'*80)
        buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()
    except ImportError: return None

html_idx = generar_html_indices()
word_idx = generar_word_indices()

st.markdown("""<div style="background:#f5f3ff;border:1px solid #7c3aed;border-radius:10px;
    padding:14px 18px;margin-bottom:16px;font-size:.87rem;color:#4c1d95;">
    📋 <strong>El documento Word es completamente editable</strong> — incluye todos los índices y espacio para tus anotaciones.
</div>""", unsafe_allow_html=True)

c1,c2,c3 = st.columns(3)
with c1:
    st.download_button("📄 Descargar Word (.docx)", data=word_idx if word_idx else b"",
        file_name=f"indices_estrategicos_{nom_sector[:15].replace(' ','_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary", use_container_width=True, disabled=word_idx is None)
with c2:
    st.download_button("🌐 Descargar HTML", data=html_idx,
        file_name=f"indices_estrategicos_{nom_sector[:15].replace(' ','_')}.html",
        mime="text/html", use_container_width=True)
with c3:
    st.info("Para PDF: abre el HTML → **Ctrl+P** → **Guardar como PDF**")
    mostrar_felix(pagina='indices')