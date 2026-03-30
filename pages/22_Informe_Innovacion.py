import streamlit as st
from asistentes import mostrar_felix, aviso_promedios
import pandas as pd
import plotly.graph_objects as go
import json, os, numpy as np, requests, unicodedata
from streamlit_echarts import st_echarts
from datetime import date

st.set_page_config(page_title="Informe de Innovación", layout="wide")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Rajdhani:wght@400;600;700&family=Inter:wght@300;400;500&display=swap');
.report-header {
    background:linear-gradient(135deg,#0a1a0e,#0d2a1a);
    border:1px solid #1a3a2a; border-radius:14px;
    padding:28px 32px; margin-bottom:24px;
}
.section-title {
    font-family:'Rajdhani',sans-serif; font-size:1.2rem; font-weight:700; color:#fff;
    background:linear-gradient(90deg,rgba(16,185,129,0.2),rgba(16,185,129,0));
    border-left:5px solid #10b981; padding:10px 16px; border-radius:0 8px 8px 0;
    margin:28px 0 14px 0; text-shadow:0 0 12px rgba(16,185,129,0.6);
}
.text-block {
    background:linear-gradient(135deg,#0a1a0e,#0d2a1a);
    border:1px solid #1a3a2a; border-radius:12px;
    padding:22px 26px; margin-bottom:16px;
    color:#cbd5e1; line-height:1.8; font-size:.93rem;
}
.rec-block {
    background:linear-gradient(135deg,#1a1200,#2a1e00);
    border-left:4px solid #f59e0b; border-radius:0 10px 10px 0;
    padding:18px 22px; margin-bottom:14px;
    color:#cbd5e1; line-height:1.8; font-size:.93rem;
}
.pred-block {
    background:linear-gradient(135deg,#0a0a2e,#1a1035);
    border-left:4px solid #7c3aed; border-radius:0 10px 10px 0;
    padding:18px 22px; margin-bottom:14px;
    color:#cbd5e1; line-height:1.8; font-size:.93rem;
}
.pending {
    color:#475569; font-style:italic; background:#0a1a0e;
    border:1px dashed #1a3a2a; border-radius:10px;
    padding:18px 22px; margin-bottom:14px;
}
.kpi-inn {
    background:linear-gradient(135deg,#0a1a0e,#0d2a1a);
    border:1px solid #1a3a2a; border-radius:12px;
    padding:14px; text-align:center;
}
</style>
""", unsafe_allow_html=True)

def cargar_api_key():
    try: return st.secrets["ANTHROPIC_API_KEY"]
    except: pass
    key = os.environ.get("ANTHROPIC_API_KEY","")
    if key: return key
    env = os.path.join(os.path.dirname(os.path.abspath(__file__)),'..', '.env')
    if os.path.exists(env):
        with open(env,'r') as f:
            for line in f:
                if line.startswith("ANTHROPIC_API_KEY"):
                    return line.split("=",1)[1].strip().strip('"').strip("'")
    return ""

ANTHROPIC_API_KEY = cargar_api_key()

PERFIL_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)),'..','perfil_empresa.json')
def cargar_perfil():
    if os.path.exists(PERFIL_FILE):
        with open(PERFIL_FILE,'r',encoding='utf-8') as f: return json.load(f)
    return {}
perfil = cargar_perfil()
for k,v in perfil.items():
    if k not in st.session_state: st.session_state[k] = v

# ── Cargar desde Supabase si falta perfil ────────────────────────────────
if not st.session_state.get('save_reg_user'):
    try:
        from supabase import create_client
        _url = st.secrets.get("SUPABASE_URL", os.environ.get("SUPABASE_URL",""))
        _key = st.secrets.get("SUPABASE_KEY", os.environ.get("SUPABASE_KEY",""))
        _ec  = st.session_state.get('empresa_codigo','') or perfil.get('sesion_codigo','')
        if _url and _key and _ec:
            _sb = create_client(_url, _key)
            _ed = _sb.table('empresas').select('*').eq('codigo',_ec).execute().data
            if _ed:
                _e = _ed[0]
                _SM = {"Alimentación y bebidas":1,"Textil y confección":2,"Cuero y calzado":3,
                    "Química y plásticos":4,"Minerales no metálicos":5,"Metalmecanico":6,
                    "Maquinaria equipo":7,"Otras manufacturas":8,"Electrónica, telecomunicaciones":9,
                    "Informática, software. Robótica, IA":10,"Actividades I+D: biotech, farmacia":11,
                    "Transporte y logística":12,"Consultoría y servicios profesionales":13,
                    "Turismo y hosteleria":14,"Retail y comercio":15,"Otros servicios":16}
                _MM = {1:1,2:1,3:1,4:1,5:1,6:1,7:1,8:1,9:2,10:2,11:2,12:3,13:3,14:3,15:3,16:3}
                _sc = _SM.get(_e.get('sector',''),0)
                st.session_state.update({
                    'save_sector_nombre':_e.get('sector',''),'save_tam_nombre':_e.get('tamano',''),
                    'save_reg_nombre':_e.get('region',''),'save_export_nombre':_e.get('exportacion',''),
                    'save_anti_nombre':_e.get('antiguedad',''),
                    'save_ventas':float(_e.get('ventas') or 0),
                    'save_empleados':int(_e.get('empleados') or 0),
                    'save_roa':float(_e.get('roa') or 0),
                    'save_var_vtas':float(_e.get('var_ventas') or 0),
                    'save_var_empl':float(_e.get('var_empleados') or 0),
                    'save_productiv':float(_e.get('productividad') or 0),
                    'save_coste_emp':float(_e.get('coste_empleado') or 0),
                    'save_endeud':float(_e.get('endeudamiento') or 0),
                    'save_reg_user':{"Andalucia":1,"Aragon":2,"Asturias":3,"Baleares":4,"Canarias":5,
                        "Cantabria":6,"Castilla la Mancha":7,"Castilla y León":8,"Cataluña":9,
                        "Com Valenciana":10,"Extremadura":11,"Galicia":12,"Madrid":13,
                        "Murcia":14,"Navarra":15,"Pais Vasco":16}.get(_e.get('region',''),0),
                    'save_tam_user':{"Pequeña":1,"Mediana":2,"Grande":3}.get(_e.get('tamano',''),0),
                    'save_sector_cod':_sc,'save_macro_cod':_MM.get(_sc,3),
                    'save_export_cod':{"Menos 10 %":1,"10 - 30 %":2,"30 - 60 %":3,"> 60 %":4}.get(_e.get('exportacion',''),0),
                    'save_anti_cod':{"Menos 10 años":1,"10-30 años":2,"> 30 años":3}.get(_e.get('antiguedad',''),0),
                    'empresa_codigo': _ec,
                })
    except Exception: pass

if not st.session_state.get('save_reg_user'):
    st.warning("⚠️ Primero completa tu perfil en **Mi Empresa → Datos de la empresa**."); st.stop()
if not st.session_state.get('score_b1'):
    try:
        from supabase import create_client as _sc
        _url = st.secrets.get("SUPABASE_URL", os.environ.get("SUPABASE_URL",""))
        _key = st.secrets.get("SUPABASE_KEY", os.environ.get("SUPABASE_KEY",""))
        _ec  = st.session_state.get('empresa_codigo','') or perfil.get('sesion_codigo','')
        if _url and _key and _ec:
            _sb = _sc(_url, _key)
            _resp = _sb.table('respuestas').select('*').eq('empresa_codigo',_ec).execute().data or []
            if _resp:
                from collections import defaultdict as _dd
                _sumas = _dd(list)
                for _r in _resp:
                    _sumas[f"B{_r['bloque']}_{_r['item']}"].append(_r['valor'])
                _proms = {k: round(sum(v)/len(v),2) for k,v in _sumas.items()}
                for _b in range(1,6):
                    _items = [v for k,v in _proms.items() if k.startswith(f"B{_b}_")]
                    if _items: st.session_state[f'score_b{_b}'] = round(sum(_items)/len(_items),2)
                if len(set(_r['bloque'] for _r in _resp)) >= 5:
                    st.session_state['informes_activados'] = True
    except Exception: pass
ruta = os.path.join(os.path.dirname(os.path.abspath(__file__)),'..','datos.xlsx')
if not os.path.exists(ruta): ruta = "datos.xlsx"
try:
    df = pd.read_excel(ruta); df.columns = df.columns.str.strip()
except Exception as e:
    st.error(f"Error: {e}"); st.stop()

EXPORTACIONES = {1:"Menos 10%",2:"10-30%",3:"30-60%",4:">60%"}
nom_sector = st.session_state.get('save_sector_nombre','—')
nom_reg    = st.session_state.get('save_reg_nombre','—')
nom_tam    = st.session_state.get('save_tam_nombre','—')
nom_export = EXPORTACIONES.get(st.session_state.get('save_export_cod',1),'—')
nom_anti   = st.session_state.get('save_anti_nombre','—')
mac_cod    = st.session_state.get('save_macro_cod',1)
tam_cod    = st.session_state.get('save_tam_user',1)
hoy        = date.today().strftime("%d/%m/%Y")

dff = df[(df['Macrosector']==mac_cod)&(df['Tamaño']==tam_cod)]
if len(dff) < 10: dff = df[df['Macrosector']==mac_cod]
if len(dff) < 10: dff = df.copy()
n_ref = len(dff)

b1 = st.session_state.get('score_b1', 0)
b2 = st.session_state.get('score_b2', 0)
b3 = st.session_state.get('score_b3', 0)
b4 = st.session_state.get('score_b4', 0)
b5 = st.session_state.get('score_b5', 0)
macro_inn = (b1*0.20 + b2*0.20 + b3*0.20 + b4*0.20 + b5*0.20) if any([b1,b2,b3,b4,b5]) else 0

SUBS = {
    'SUB_1.1_Dpto':     ('Dpto. y Recursos I+D',       'score_sub1_1'),
    'SUB_1.2_PresupID': ('Presupuesto I+D',             'score_sub1_2'),
    'SUB_1.3_Gasto':    ('Gasto en Innovación',         'score_sub1_3'),
    'SUB_2.1_Gbasica':  ('Gestión Básica Proy.',        'score_sub2_1'),
    'SUB_2.2_Gavanz':   ('Gestión Avanzada Proy.',      'score_sub2_2'),
    'SUB_2.3_OrgProy':  ('Organización Proyectos',      'score_sub2_3'),
    'SUB_2.4_EvalRdto': ('Evaluación Rendimiento',      'score_sub2_4'),
    'SUB_3.1_EstrDes':  ('Estrategia Nuevos Prod.',     'score_sub3_1'),
    'SUB_3.2_OportMdo': ('Oportunidad de Mercado',      'score_sub3_2'),
    'SUB_3.3_RecepNprod':('Receptividad Cliente',       'score_sub3_3'),
    'SUB_3.4_Clientes': ('Orientación al Cliente',      'score_sub3_4'),
    'SUB_3.5_ImpacNprod':('Viabilidad del Producto',    'score_sub3_5'),
    'SUB_4.1_InnEstrat':('Inn. Estratégica',            'score_sub4_1'),
    'SUB_4.2_CultInn':  ('Cultura Innovadora',          'score_sub4_2'),
    'SUB_4.3_Obstac':   ('Gestión Obstáculos',          'score_sub4_3'),
    'SUB_4.4_InnAb':    ('Innovación Abierta',          'score_sub4_4'),
    'SUB_4.5_Creativ':  ('Creatividad y Talento',       'score_sub4_5'),
    'SUB_5.1_ImpEstim': ('Impacto Estimado',            'score_sub5_1'),
    'SUB_5.2_InnEfect': ('Efecto Real Innovación',      'score_sub5_2'),
}

IND_INFO = {
    'IND_IDi':      ('I+D+i',               b1, ['SUB_1.1_Dpto','SUB_1.2_PresupID','SUB_1.3_Gasto']),
    'IND_GPROY':    ('Gestión Proyectos',    b2, ['SUB_2.1_Gbasica','SUB_2.2_Gavanz','SUB_2.3_OrgProy','SUB_2.4_EvalRdto']),
    'IND_DESPROD':  ('Desarrollo Productos', b3, ['SUB_3.1_EstrDes','SUB_3.2_OportMdo','SUB_3.3_RecepNprod','SUB_3.4_Clientes','SUB_3.5_ImpacNprod']),
    'IND_ESTRINN':  ('Estrategia Innovación',b4, ['SUB_4.1_InnEstrat','SUB_4.2_CultInn','SUB_4.3_Obstac','SUB_4.4_InnAb','SUB_4.5_Creativ']),
    'IND_DESMPINN': ('Desempeño Innovación', b5, ['SUB_5.1_ImpEstim','SUB_5.2_InnEfect']),
}

sub_vals = {col: st.session_state.get(key, 0) for col, (nombre, key) in SUBS.items()}
ind_medias = {cod: round(dff[cod].mean(),2) if cod in dff.columns else 3.0 for cod in IND_INFO}
sub_medias = {col: round(dff[col].mean(),2) if col in dff.columns else 3.0 for col in SUBS}
macro_media = round(dff['MACRO_INNOVACION'].mean(),2) if 'MACRO_INNOVACION' in dff.columns else 3.0

def pct(series, val, inv=False):
    if inv: return float(np.sum(series >= val)/len(series)*100)
    return float(np.sum(series <= val)/len(series)*100)

def col_v(v, mx=5):
    r = v/mx
    if r >= 0.66: return '#10b981'
    if r >= 0.40: return '#f59e0b'
    return '#ef4444'

def col_pct(v):
    if v >= 66: return '#10b981'
    if v >= 33: return '#f59e0b'
    return '#ef4444'

def nivel(v, mx=5):
    r = v/mx
    if r >= 0.66: return 'Alto'
    if r >= 0.40: return 'Medio'
    return 'Bajo'

def norm(s):
    return ''.join(c for c in unicodedata.normalize('NFD',s) if unicodedata.category(c) != 'Mn').upper()

ind_pcts = {cod: round(pct(dff[cod], vals[1])) if cod in dff.columns else 50 for cod, vals in IND_INFO.items()}
macro_pct = round(pct(dff['MACRO_INNOVACION'], macro_inn)) if 'MACRO_INNOVACION' in dff.columns else 50

def resumen_brechas():
    lineas = []
    for cod, (nombre, val, subs_lista) in IND_INFO.items():
        media = ind_medias[cod]; p_ind = ind_pcts[cod]
        lineas.append(f"\n[{nombre}]: {val:.2f}/5 vs media {media:.2f} (P{p_ind})")
        for sub in subs_lista:
            sv = sub_vals.get(sub, 0); sm = sub_medias.get(sub, 3.0); sn = SUBS[sub][0]
            if sv > 0:
                lineas.append(f"  • {sn}: {sv:.2f}/5 vs {sm:.2f} (gap {sv-sm:+.2f})")
    return "\n".join(lineas)

def llamar_ia():
    if not ANTHROPIC_API_KEY: return "ERROR: API Key no configurada."
    brechas = resumen_brechas()
    ind_resumen = " | ".join([f"{vals[0]}: {vals[1]:.2f}/5 (P{ind_pcts[cod]}°)" for cod, vals in IND_INFO.items()])
    prompt = f"""Eres un consultor experto en innovación empresarial. Redacta un informe ejecutivo de diagnóstico de innovación en español. Máximo 750 palabras. Tono directo y profesional. Escribe en párrafos narrativos fluidos.

PERFIL: {nom_sector} | {nom_tam} | {nom_reg} | Exportación: {nom_export} | Antigüedad: {nom_anti}
Grupo de referencia: {n_ref} empresas comparables.
ÍNDICE GLOBAL DE INNOVACIÓN: {macro_inn:.2f}/5 (percentil {macro_pct}° en el grupo)
LOS 5 INDICADORES: {ind_resumen}
ANÁLISIS DETALLADO: {brechas}

REGLAS DE FORMATO: Cada vez que menciones un indicador escríbelo entre corchetes dobles: [[I+D+i]], [[Gestión Proyectos]], [[Desarrollo Productos]], [[Estrategia Innovación]], [[Desempeño Innovación]].

TONO: Consultor senior estilo McKinsey. PROHIBIDO: "alarmante", "crítico", "catastrófico", "grave", "urgente". USA: "por debajo de la media", "margen de mejora", "área prioritaria", "posición sólida".

ESTRUCTURA EXACTA con estos 3 títulos con ##:
## DIAGNOSTICO DE INNOVACION
## FORTALEZAS Y BRECHAS
## RECOMENDACIONES Y PREDICCION

Escribe los títulos SIN tildes tal como aparecen arriba."""
    try:
        resp = requests.post("https://api.anthropic.com/v1/messages",
            headers={"Content-Type":"application/json","x-api-key":ANTHROPIC_API_KEY,"anthropic-version":"2023-06-01"},
            json={"model":"claude-sonnet-4-20250514","max_tokens":1600,"messages":[{"role":"user","content":prompt}]},
            timeout=45)
        data = resp.json()
        if 'content' in data and data['content']: return data['content'][0]['text']
        return f"ERROR: {data.get('error',{}).get('message','Respuesta inesperada')}"
    except requests.exceptions.Timeout: return "ERROR: Tiempo de espera agotado."
    except Exception as e: return f"ERROR: {e}"

def parsear(texto):
    resultado = []
    for parte in [p.strip() for p in texto.split('##') if p.strip()]:
        lineas = parte.split('\n',1)
        titulo_orig = lineas[0].strip()
        cuerpo = lineas[1].strip() if len(lineas)>1 else ''
        resultado.append((norm(titulo_orig), titulo_orig, cuerpo))
    return resultado

def buscar(secciones, *kws):
    for tn, to, cu in secciones:
        for kw in kws:
            if norm(kw) in tn: return cu
    return None

def render(texto, estilo='normal'):
    if not texto: return ''
    import re
    IND_COLORES_TEXT = {
        'I+D+i':'#3b82f6','Gestión Proyectos':'#8b5cf6','Gestión de Proyectos':'#8b5cf6',
        'Desarrollo Productos':'#10b981','Estrategia Innovación':'#f59e0b',
        'Estrategia de Innovación':'#f59e0b','Desempeño Innovación':'#ef4444',
        'Desempeño de la Innovación':'#ef4444',
    }
    def reemplazar(m):
        nombre = m.group(1).strip()
        color = '#00d4ff'
        for k,c in IND_COLORES_TEXT.items():
            if k.lower() in nombre.lower(): color=c; break
        return f'<span style="color:{color};font-weight:700;background:rgba(255,255,255,0.06);border-radius:4px;padding:1px 6px;">{nombre}</span>'
    texto = re.sub(r'\[\[(.+?)\]\]', reemplazar, texto)
    html = texto.replace('\n\n','</p><p>').replace('\n','<br>')
    css = {'normal':'text-block','rec':'rec-block','pred':'pred-block'}.get(estilo,'text-block')
    return f'<div class="{css}"><p>{html}</p></div>'

# ── Cabecera ──────────────────────────────────────────────────────────────
mostrar_felix(pagina='informe_innovacion')
aviso_promedios()
macro_c = col_v(macro_inn)
st.markdown(f"""
<div class="report-header">
  <div style="display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:16px;">
    <div>
      <div style="font-size:.7rem;color:#64748b;letter-spacing:2px;text-transform:uppercase;margin-bottom:4px;">
        Informe de Innovación · Diagnóstico por Bloques</div>
      <h1 style="font-family:Rajdhani,sans-serif;color:#10b981;margin:0;font-size:2rem;">{nom_sector}</h1>
      <p style="color:#94a3b8;margin:6px 0 0 0;font-size:.88rem;">
        {nom_tam} · {nom_reg} · Exportación: {nom_export} · {nom_anti}<br>
        <span style="color:#64748b;font-size:.8rem;">Referencia: {n_ref} empresas · {hoy}</span>
      </p>
    </div>
    <div style="text-align:center;background:rgba(0,0,0,0.35);border:1px solid {macro_c}55;border-radius:12px;padding:16px 28px;">
      <div style="font-size:.65rem;color:#64748b;letter-spacing:1.5px;text-transform:uppercase;">Índice Global de Innovación</div>
      <div style="font-family:Rajdhani,sans-serif;font-size:4.2rem;font-weight:700;color:{macro_c};line-height:1;">{macro_inn:.2f}</div>
      <div style="font-size:.78rem;color:#94a3b8;">/ 5 · Percentil {macro_pct}° · {nivel(macro_inn)}</div>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)

kcols = st.columns(5)
for col_st, (cod, (nombre, val, _)) in zip(kcols, IND_INFO.items()):
    p_ind = ind_pcts[cod]
    col_st.markdown(f"""<div class="kpi-inn">
      <div style="font-size:.62rem;color:#64748b;">{nombre}</div>
      <div style="font-family:Rajdhani,sans-serif;font-size:1.9rem;font-weight:700;color:{col_v(val)};">{val:.2f}</div>
      <div style="font-size:.7rem;color:#64748b;">/ 5 · P{p_ind}°</div>
    </div>""", unsafe_allow_html=True)

# ── Velocímetros ──────────────────────────────────────────────────────────
st.markdown('<div class="section-title">⚡ Los 5 Indicadores de Innovación</div>', unsafe_allow_html=True)
gcols = st.columns(5)
for col_st, (cod, (nombre, val, _)) in zip(gcols, IND_INFO.items()):
    media = ind_medias[cod]
    fig = go.Figure(go.Indicator(
        mode="gauge+number", value=val,
        number={'font':{'size':26,'color':col_v(val)},'valueformat':'.2f'},
        title={'text':f"{nombre[:18]}<br><span style='font-size:8px;color:#64748b;'>Media: {media:.2f}</span>",'font':{'size':10,'color':'#94a3b8'}},
        gauge={'axis':{'range':[0,5],'tickcolor':'#334155','tickfont':{'color':'#64748b','size':7}},
            'bar':{'color':col_v(val),'thickness':0.28},'bgcolor':'#0a0e1a','borderwidth':0,
            'steps':[{'range':[0,2],'color':'#1a0505'},{'range':[2,3.3],'color':'#1a1200'},{'range':[3.3,5],'color':'#051a0a'}],
            'threshold':{'line':{'color':'#475569','width':2},'thickness':0.75,'value':media}}
    ))
    fig.update_layout(height=195,margin=dict(l=8,r=8,t=45,b=4),paper_bgcolor='rgba(0,0,0,0)')
    col_st.plotly_chart(fig, use_container_width=True)

# ── Barras agrupadas ──────────────────────────────────────────────────────
st.markdown('<div class="section-title">📊 Indicadores y Subindicadores vs. Grupo</div>', unsafe_allow_html=True)
IND_COLORES = {'IND_IDi':'#3b82f6','IND_GPROY':'#8b5cf6','IND_DESPROD':'#10b981','IND_ESTRINN':'#f59e0b','IND_DESMPINN':'#ef4444'}
SUB_COLORES = {'IND_IDi':'rgba(59,130,246,0.55)','IND_GPROY':'rgba(139,92,246,0.55)','IND_DESPROD':'rgba(16,185,129,0.55)','IND_ESTRINN':'rgba(245,158,11,0.55)','IND_DESMPINN':'rgba(239,68,68,0.55)'}
categorias=[]; vals_emp=[]; vals_grp=[]; colores_emp=[]; es_ind=[]
for cod, (nombre, val, subs_lista) in IND_INFO.items():
    categorias.append(nombre[:16]); vals_emp.append(round(val,2)); vals_grp.append(ind_medias[cod])
    colores_emp.append(IND_COLORES[cod]); es_ind.append(True)
    for sub in subs_lista:
        sv=sub_vals.get(sub,0); sm=sub_medias.get(sub,3.0); sn=SUBS[sub][0][:20]
        categorias.append(f"  ↳ {sn}"); vals_emp.append(round(sv,2) if sv>0 else 0)
        vals_grp.append(sm); colores_emp.append(SUB_COLORES[cod]); es_ind.append(False)
bar_h = max(500, len(categorias)*26)
bar_opt = {
    "backgroundColor":"#0a1a0e","tooltip":{"trigger":"axis","axisPointer":{"type":"shadow"},"backgroundColor":"#1e293b","textStyle":{"color":"#e2e8f0"}},
    "legend":{"data":["Mi Empresa","Media Grupo"],"textStyle":{"color":"#94a3b8"},"top":5},
    "grid":{"left":"2%","right":"8%","bottom":"3%","top":"8%","containLabel":True},
    "xAxis":{"type":"value","min":0,"max":5,"axisLabel":{"color":"#64748b"},"splitLine":{"lineStyle":{"color":"#1e3a2a"}}},
    "yAxis":{"type":"category","data":categorias,"inverse":True,"axisLabel":{"color":"#94a3b8","fontSize":10}},
    "series":[
        {"name":"Mi Empresa","type":"bar","data":[{"value":v,"itemStyle":{"color":c,"borderRadius":[0,4,4,0]}} for v,c in zip(vals_emp,colores_emp)],"barMaxWidth":18},
        {"name":"Media Grupo","type":"bar","data":[{"value":v,"itemStyle":{"color":"rgba(100,116,139,0.4)","borderRadius":[0,4,4,0]}} for v in vals_grp],"barMaxWidth":10},
    ],"barGap":"10%",
}
st_echarts(options=bar_opt, height=f"{bar_h}px")

# ── Botón IA ──────────────────────────────────────────────────────────────
st.divider()
col_btn, col_hint = st.columns([1,2])
with col_btn:
    generar = st.button("🚀 Generar Informe de Innovación con IA", type="primary", use_container_width=True)
with col_hint:
    st.info("La IA analiza todos los bloques y subindicadores. ~15-25 segundos.")
if generar:
    with st.spinner("Analizando tu perfil innovador..."):
        st.session_state['inf2_texto'] = llamar_ia()

texto_raw = st.session_state.get('inf2_texto','')
if texto_raw.startswith('ERROR:'):
    st.error(f"❌ {texto_raw}"); st.stop()
secciones = parsear(texto_raw) if texto_raw else []

# ── Secciones IA ─────────────────────────────────────────────────────────
st.markdown('<div class="section-title">🔍 Diagnóstico de Innovación</div>', unsafe_allow_html=True)
txt = buscar(secciones,'DIAGNOSTICO','INNOVACION','PERFIL')
st.markdown(render(txt) if txt else '<div class="pending">Pulsa "Generar Informe" para ver el diagnóstico.</div>', unsafe_allow_html=True)

st.markdown('<div class="section-title">🕸️ Radar de los 5 Indicadores vs. Grupo</div>', unsafe_allow_html=True)
ind_labels=[v[0] for v in IND_INFO.values()]; emp_ind=[round(v[1],2) for v in IND_INFO.values()]; grp_ind=[ind_medias[c] for c in IND_INFO]
cr1,_ = st.columns([1,1])
with cr1:
    r1_opt = {
        "backgroundColor":"#0a1a0e","tooltip":{"trigger":"item","backgroundColor":"#1e293b","textStyle":{"color":"#e2e8f0"}},
        "legend":{"data":["Mi Empresa","Media Grupo"],"textStyle":{"color":"#94a3b8"},"bottom":0},
        "radar":{"indicator":[{"name":n,"max":5} for n in ind_labels],"shape":"polygon","splitNumber":5,
            "axisName":{"color":"#94a3b8","fontSize":13},
            "splitLine":{"lineStyle":{"color":["#1e3a2a"]*5}},
            "splitArea":{"areaStyle":{"color":["rgba(16,185,129,0.02)","rgba(16,185,129,0.06)"]}},
            "axisLine":{"lineStyle":{"color":"#1e3a2a"}},"radius":"65%"},
        "series":[{"type":"radar","data":[
            {"value":emp_ind,"name":"Mi Empresa","itemStyle":{"color":"#10b981"},"lineStyle":{"color":"#10b981","width":2},"areaStyle":{"color":"rgba(16,185,129,0.2)"},"symbol":"circle","symbolSize":7},
            {"value":grp_ind,"name":"Media Grupo","itemStyle":{"color":"#f59e0b"},"lineStyle":{"color":"#f59e0b","width":1.5,"type":"dashed"},"areaStyle":{"color":"rgba(245,158,11,0.06)"}},
        ]}],
    }
    st_echarts(options=r1_opt, height="400px")

st.markdown('<div class="section-title">⚖️ Fortalezas y Brechas de Innovación</div>', unsafe_allow_html=True)
cf, crank = st.columns([1,1])
with cf:
    txt = buscar(secciones,'FORTALEZA','BRECHA','DEBILIDAD')
    st.markdown(render(txt) if txt else '<div class="pending">Fortalezas y brechas pendientes.</div>', unsafe_allow_html=True)
with crank:
    st.caption("**Ranking Percentil de Subindicadores**")
    sub_pcts_lista = [(SUBS[s][0][:25], round(pct(dff[s], sub_vals.get(s,0))))
        for s in SUBS if sub_vals.get(s,0)>0 and s in dff.columns]
    sub_pcts_lista.sort(key=lambda x: x[1])
    if sub_pcts_lista:
        s_labels=[x[0] for x in sub_pcts_lista]; s_vals=[x[1] for x in sub_pcts_lista]
        rank_sub_opt = {
            "backgroundColor":"#0a1a0e","tooltip":{"trigger":"axis","backgroundColor":"#1e293b","textStyle":{"color":"#e2e8f0"}},
            "grid":{"left":"2%","right":"12%","bottom":"3%","containLabel":True},
            "xAxis":{"type":"value","min":0,"max":100,"axisLabel":{"color":"#64748b"},"splitLine":{"lineStyle":{"color":"#1e3a2a"}}},
            "yAxis":{"type":"category","data":s_labels,"inverse":True,"axisLabel":{"color":"#94a3b8","fontSize":9}},
            "series":[{"type":"bar",
                "data":[{"value":v,"itemStyle":{"color":col_pct(v),"borderRadius":[0,4,4,0]}} for v in s_vals],
                "label":{"show":True,"position":"right","color":"#e2e8f0","fontSize":10,"formatter":"{c}°"},
                "markLine":{"silent":True,"lineStyle":{"color":"#f59e0b","type":"dashed","width":1.5},
                    "data":[{"xAxis":50,"label":{"formatter":"Mediana","color":"#f59e0b","fontSize":9}}]},
            }]
        }
        st_echarts(options=rank_sub_opt, height=f"{max(350,len(s_labels)*28)}px")

st.markdown('<div class="section-title">🎯 Recomendaciones y Predicción</div>', unsafe_allow_html=True)
txt_rec = buscar(secciones,'RECOMENDACION','PREDICCION','ACCION')
if txt_rec:
    parrafos=[p.strip() for p in txt_rec.split('\n\n') if p.strip()]
    recs=[p for p in parrafos if p and p[0].isdigit() and '.' in p[:3]]
    preds=[p for p in parrafos if p not in recs]
    if not recs and not preds: st.markdown(render(txt_rec,'rec'), unsafe_allow_html=True)
    else:
        if recs: st.markdown(f'<div class="rec-block"><p>{"</p><p>".join(recs)}</p></div>', unsafe_allow_html=True)
        if preds: st.markdown(f'<div class="pred-block"><p>{"</p><p>".join(preds)}</p></div>', unsafe_allow_html=True)
else:
    c1,c2=st.columns(2)
    c1.markdown('<div class="pending">Recomendaciones pendientes.</div>', unsafe_allow_html=True)
    c2.markdown('<div class="pending">Predicción pendiente.</div>', unsafe_allow_html=True)

st.divider()

# ── Descarga ──────────────────────────────────────────────────────────────
st.markdown('<div class="section-title">📥 Descargar Informe</div>', unsafe_allow_html=True)

def generar_html():
    ia_html = ""
    colores = {'DIAGNOSTICO':'#0a7a3a','FORTALEZA':'#dc2626','BRECHA':'#dc2626','RECOMENDACION':'#d97706','PREDICCION':'#7c3aed'}
    for tn,to,cu in secciones:
        color='#0a7a3a'
        for kw,c in colores.items():
            if kw in tn: color=c; break
        ia_html += f"<h2 style='color:{color};border-left:4px solid {color};padding-left:10px;margin-top:28px;'>{to}</h2><p style='line-height:1.75;'>{cu.replace(chr(10),'<br>')}</p>"
    ind_tabla = "".join([
        f"<tr><td><b>{cod}</b></td><td>{vals[0]}</td><td style='text-align:center;font-weight:700;color:{'green' if vals[1]>=3.3 else 'orange' if vals[1]>=2 else 'red'};'>{vals[1]:.2f}/5</td><td style='text-align:center;color:#666;'>{ind_medias[cod]:.2f}/5</td><td style='text-align:center;'>{ind_pcts[cod]}°</td></tr>"
        for cod,vals in IND_INFO.items()])
    sub_tabla = "".join([
        f"<tr><td>{SUBS[s][0]}</td><td style='text-align:center;'>{sub_vals.get(s,0):.2f}/5</td><td style='text-align:center;color:#666;'>{sub_medias.get(s,3):.2f}/5</td></tr>"
        for s in SUBS if sub_vals.get(s,0)>0])
    inn_color = 'green' if macro_inn>=3.3 else 'orange' if macro_inn>=2 else 'red'
    return f"""<!DOCTYPE html><html><head><meta charset='utf-8'>
<style>body{{font-family:Georgia,serif;margin:50px auto;max-width:860px;color:#1a1a1a;line-height:1.75;}}
h1{{color:#0a7a3a;border-bottom:3px solid #0a7a3a;padding-bottom:10px;font-size:1.8rem;}}
.perfil{{background:#f0fff4;border-radius:8px;padding:16px;margin:16px 0;font-size:.9rem;}}
.inn{{font-size:2.8rem;font-weight:700;color:{inn_color};}}
table{{border-collapse:collapse;width:100%;margin:16px 0;font-size:.9rem;}}
th{{background:#0a7a3a;color:white;padding:10px;text-align:left;}}
td{{padding:9px 11px;border-bottom:1px solid #e5e7eb;}}
.notas{{border:1px dashed #9ca3af;border-radius:8px;padding:40px 16px;margin:16px 0;color:#9ca3af;font-style:italic;text-align:center;}}
@media print{{body{{margin:20px;}}}}</style></head><body>
<h1>Informe de Innovación · Diagnóstico por Bloques</h1>
<div class="perfil"><strong>{nom_sector}</strong> | {nom_tam} | {nom_reg} | Exportación: {nom_export} | {nom_anti}<br>
Referencia: <strong>{n_ref} empresas comparables</strong> | Fecha: {hoy}</div>
<p>Índice Global de Innovación: <span class="inn">{macro_inn:.2f}</span>/5 · Percentil {macro_pct}° · {nivel(macro_inn)}</p>
<h2 style='color:#0a7a3a;'>Indicadores de Innovación</h2>
<table><tr><th>Código</th><th>Indicador</th><th>Mi Empresa</th><th>Media Grupo</th><th>Percentil</th></tr>{ind_tabla}</table>
<h2 style='color:#0a7a3a;'>Subindicadores</h2>
<table><tr><th>Subindicador</th><th>Mi Empresa</th><th>Media Grupo</th></tr>{sub_tabla}</table>
{ia_html or '<p style="color:#6b7280;font-style:italic;">Genera el análisis IA antes de descargar.</p>'}
<div class="notas">✏️ Espacio para anotaciones y comentarios del equipo directivo</div>
<hr style='margin-top:40px;'/>
<p style='color:#9ca3af;font-size:.78rem;text-align:center;'>Plataforma de Diagnóstico Estratégico · {hoy}</p>
</body></html>"""

def generar_word():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        doc = Document()
        t = doc.add_heading('Informe de Innovación · Diagnóstico por Bloques', 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Empresa: {nom_sector} | {nom_tam} | {nom_reg} | Exportación: {nom_export} | {nom_anti}")
        doc.add_paragraph(f"Referencia: {n_ref} empresas | Fecha: {hoy}")
        doc.add_paragraph(f"Índice Global de Innovación: {macro_inn:.2f}/5 — Percentil {macro_pct}° — {nivel(macro_inn)}")
        doc.add_heading('Indicadores de Innovación', level=1)
        tabla=doc.add_table(rows=1,cols=4); tabla.style='Table Grid'
        hdr=tabla.rows[0].cells; hdr[0].text='Indicador'; hdr[1].text='Mi Empresa'; hdr[2].text='Media Grupo'; hdr[3].text='Percentil'
        for cod,(nombre,val,_) in IND_INFO.items():
            row=tabla.add_row().cells; row[0].text=nombre; row[1].text=f"{val:.2f}/5"
            row[2].text=f"{ind_medias[cod]:.2f}/5"; row[3].text=f"{ind_pcts[cod]}°"
        doc.add_heading('Subindicadores', level=1)
        tabla2=doc.add_table(rows=1,cols=3); tabla2.style='Table Grid'
        hdr2=tabla2.rows[0].cells; hdr2[0].text='Subindicador'; hdr2[1].text='Mi Empresa'; hdr2[2].text='Media Grupo'
        for s,(sn,_) in SUBS.items():
            if sub_vals.get(s,0)>0:
                row=tabla2.add_row().cells; row[0].text=sn; row[1].text=f"{sub_vals[s]:.2f}/5"; row[2].text=f"{sub_medias.get(s,3):.2f}/5"
        if secciones:
            for tn,to,cuerpo in secciones:
                doc.add_heading(to,level=1); doc.add_paragraph(cuerpo)
        else:
            doc.add_paragraph("Genera el análisis IA antes de descargar.")
        doc.add_heading('Notas y comentarios del equipo directivo', level=1)
        doc.add_paragraph("[ Escribe aquí tus conclusiones y próximos pasos ]")
        for _ in range(8):
            p=doc.add_paragraph(); p.add_run('_'*80)
        buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()
    except ImportError: return None

html_out=generar_html(); word_out=generar_word()
st.markdown("""<div style="background:#f0fdf4;border:1px solid #10b981;border-radius:10px;
    padding:14px 18px;margin-bottom:16px;font-size:.87rem;color:#065f46;">
    📋 <strong>El documento Word es completamente editable</strong> — incluye indicadores, subindicadores y el análisis IA.
</div>""", unsafe_allow_html=True)
c1,c2,c3=st.columns(3)
with c1:
    st.download_button("📄 Descargar Word (.docx)", data=word_out if word_out else b"",
        file_name=f"informe_innovacion_{nom_sector[:15].replace(' ','_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary", use_container_width=True, disabled=word_out is None)
with c2:
    st.download_button("🌐 Descargar HTML", data=html_out,
        file_name=f"informe_innovacion_{nom_sector[:15].replace(' ','_')}.html",
        mime="text/html", use_container_width=True)
with c3:
    st.info("Para PDF: abre el HTML → **Ctrl+P** → **Guardar como PDF**")